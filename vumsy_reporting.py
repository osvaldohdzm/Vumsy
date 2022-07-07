import argparse
import base64
import docx
import html_to_json
import inflect
import json
import locale
import logging
import math
import mistune
import numpy as np
import os
import os.path
import pdb
import pyfiglet
import pythoncom
import re
import requests
import shutil
import sys
import unidecode
import win32clipboard as clip
import win32com
import win32com.client as win32
import xmltojson
from datetime import date, datetime, time, timedelta
from win32com.client import constants
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor
from num2words import num2words
from os.path import abspath
import locale

locale.setlocale(locale.LC_TIME, 'es_MX.UTF-8')

# Program constants definition
constants ={} 
constants["title-style"] = "Título 1" # Dependes language english or spanish

class bcolors:
    HEADER = '\033[95m'
    OKBLUE = '\033[94m'
    OKCYAN = '\033[96m'
    OKGREEN = '\033[92m'
    WARNING = '\033[93m'
    FAIL = '\033[91m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'

def num2ordinalabrev(N):
  word = num2ordinalapocade(N)
  return word[len(word) - 2:]

def num2ordinalapocade(N):
  Unidad=["", "primer", "segundo", "tercer","cuarto", "quinto", "sexto", "septimo", "octavo", "noveno"]
  Decena=["", "decimo", "vigesimo", "trigesimo", "cuadragesimo","quincuagesimo", "sexagesimo", "septuagesimo", "octogesimo", "nonagesimo"]
  Centena=["", "centesimo", "ducentesimo", "tricentesimo", " cuadringentesimo", " quingentesimo", " sexcentesimo", " septingentesimo"," octingentesimo", " noningentesimo"]
  u= N % 10
  d=int(math.floor(N/10))%10
  c=int(math.floor(N/100))
  if(N>=100): 
   return(Centena[c]+" "+Decena[d]+" "+Unidad[u])
  else:
   if(N>=10):
    return(Decena[d]+" "+Unidad[u])
   else:
    return(Unidad[N])

def format_range_remediation_status(crange,wb):
   selection = wb.Worksheets(1).Range(crange)
   # Add condition
   selection.FormatConditions.Add(Type=win32.constants.xlCellValue, Operator =win32.constants.xlEqual, Formula1 ="=\"REMEDIADA\"")
   selection.FormatConditions(selection.FormatConditions.Count).SetFirstPriority()
   selection.FormatConditions(1).Interior.PatternColorIndex = win32.constants.xlAutomatic
   selection.FormatConditions(1).Interior.ThemeColor = win32.constants.xlThemeColorAccent6
   selection.FormatConditions(1).Interior.TintAndShade = 0.59996337778862
   # Add condition
   selection = wb.Worksheets(1).Range(crange)
   selection.FormatConditions.Add(Type=win32.constants.xlCellValue, Operator =win32.constants.xlEqual, Formula1 ="=\"PARCIALMENTE REMEDIADA\"")
   selection.FormatConditions(selection.FormatConditions.Count).SetFirstPriority()
   selection.FormatConditions(1).Interior.PatternColorIndex = win32.constants.xlAutomatic
   selection.FormatConditions(1).Interior.ThemeColor = win32.constants.xlThemeColorAccent4
   selection.FormatConditions(1).Interior.TintAndShade = 0.599963377788629
   # Add condition
   selection = wb.Worksheets(1).Range(crange)
   selection.FormatConditions.Add(Type=win32.constants.xlCellValue, Operator =win32.constants.xlEqual, Formula1 ="=\"NO REMEDIADA\"")
   selection.FormatConditions(selection.FormatConditions.Count).SetFirstPriority()
   selection.FormatConditions(1).Interior.PatternColorIndex = win32.constants.xlAutomatic
   selection.FormatConditions(1).Interior.ThemeColor = win32.constants.xlThemeColorAccent2
   selection.FormatConditions(1).Interior.TintAndShade = 0.399945066682943
   
def format_range_risk_leve(crange,wb):
   # Add condition
   selection = wb.Worksheets(1).Range(crange)
   selection.FormatConditions.Add(Type=win32.constants.xlCellValue, Operator =win32.constants.xlEqual, Formula1 ="=\"CRÍTICA\"")
   selection.FormatConditions(selection.FormatConditions.Count).SetFirstPriority()
   selection.FormatConditions(1).Interior.PatternColorIndex = win32.constants.xlAutomatic
   selection.FormatConditions(1).Interior.Color = 192
   selection.FormatConditions(1).Font.ThemeColor = win32.constants.xlThemeColorDark1
   selection.FormatConditions(1).Font.TintAndShade = 0
   # Add condition
   selection = wb.Worksheets(1).Range(crange)
   selection.FormatConditions.Add(Type=win32.constants.xlCellValue, Operator =win32.constants.xlEqual, Formula1 ="=\"ALTA\"")
   selection.FormatConditions(selection.FormatConditions.Count).SetFirstPriority()
   selection.FormatConditions(1).Interior.PatternColorIndex = win32.constants.xlAutomatic
   selection.FormatConditions(1).Interior.Color = 255
   selection.FormatConditions(1).Font.ThemeColor = win32.constants.xlThemeColorDark1
   selection.FormatConditions(1).Font.TintAndShade = 0
   # Add condition
   selection = wb.Worksheets(1).Range(crange)
   selection.FormatConditions.Add(Type=win32.constants.xlCellValue, Operator =win32.constants.xlEqual, Formula1 ="=\"MEDIA\"")
   selection.FormatConditions(selection.FormatConditions.Count).SetFirstPriority()
   selection.FormatConditions(1).Interior.PatternColorIndex = win32.constants.xlAutomatic
   selection.FormatConditions(1).Interior.Color = 49407
   selection.FormatConditions(1).Font.ThemeColor = win32.constants.xlThemeColorDark1
   selection.FormatConditions(1).Font.TintAndShade = 0
   selection.FormatConditions(1).Font.ColorIndex = win32.constants.xlAutomatic
   # Add condition
   selection = wb.Worksheets(1).Range(crange)
   selection.FormatConditions.Add(Type=win32.constants.xlCellValue, Operator =win32.constants.xlEqual, Formula1 ="=\"BAJA\"")
   selection.FormatConditions(selection.FormatConditions.Count).SetFirstPriority()
   selection.FormatConditions(1).Interior.PatternColorIndex = win32.constants.xlAutomatic
   selection.FormatConditions(1).Interior.Color = 65535
   selection.FormatConditions(1).Font.TintAndShade = 0
   selection.FormatConditions(1).Font.ColorIndex = win32.constants.xlAutomatic
   # Add condition
   selection = wb.Worksheets(1).Range(crange)
   selection.FormatConditions.Add(Type=win32.constants.xlCellValue, Operator =win32.constants.xlEqual, Formula1 ="=\"INFORMATIVA\"")
   selection.FormatConditions(selection.FormatConditions.Count).SetFirstPriority()
   selection.FormatConditions(1).Interior.PatternColorIndex = win32.constants.xlAutomatic
   selection.FormatConditions(1).Interior.Color = 11382189
   selection.FormatConditions(1).Font.ThemeColor = win32.constants.xlThemeColorLight1
   selection.FormatConditions(1).Font.TintAndShade = 0
   selection.FormatConditions(1).Font.ColorIndex = win32.constants.xlAutomatic
   
def get_risk_level_from_score(risk_score):
  level = 'OTRO'
  if (float(risk_score) >= 9.0):
      level = 'CRÍTICA'
  elif (float(risk_score) >= 7.0):
      level = 'ALTA'
  elif (float(risk_score) >= 4.0):
      level = 'MEDIA'
  elif (float(risk_score) >= 0.1):
      level = 'BAJA'
  elif (float(risk_score) <= 0.0):
      level = 'INFORMATIVA'
  return level

def add_qa_vulns_to_tablefile(doct_path,wordapp, qa_data):
  if qa_data:
      doct = wordapp.Documents.Open(doct_path)
      doct.Activate()
      wordapp.Selection.EndKey(Unit=win32.constants.wdStory)
      wordapp.Selection.TypeParagraph
      wordapp.Selection.TypeText(Text='Vulnerabilidades asociadas al ambiente QA') 
      wordapp.Selection.HomeKey(Unit=win32.constants.wdLine, Extend=win32.constants.wdExtend)
      wordapp.Selection.Style = doct.Styles(constants["title-style"])
      wordapp.Selection.EndKey(Unit=win32.constants.wdLine)
      wordapp.Selection.TypeText(Text='\r') 
      wordapp.Selection.Style = wordapp.ActiveDocument.Styles("Normal")
      wordapp.Selection.TypeText(Text='<<qa_vulnerabilities_list>>\r') 
      qa_vuln_list = []
      try:
          for item in qa_data:
              qa_vuln_list.append(item) 
          wordapp.Selection.HomeKey(Unit=win32.constants.wdStory)
          wordapp.Selection.Find.Execute('<<qa_vulnerabilities_list>>') 
          wordapp.Selection.Range.ListFormat.ApplyListTemplateWithLevel(ListTemplate = wordapp.ListGalleries(win32.constants.wdBulletGallery).ListTemplates(1), ContinuePreviousList= True, ApplyTo = win32.constants.wdListApplyToWholeList, DefaultListBehavior= win32.constants.wdWord10ListBehavior)
          wordapp.Selection.Font.Name = "Montserrat"
          wordapp.Selection.Text = '\n'.join(qa_vuln_list)
      except Exception as e: 
          print(e) 
      doct.SaveAs(doct_path)
      doct.Close(False)

def create_vuln_matrix(vulns_list, output_folder, file_name_no_ext):
    if vulns_list:
        try:
            vulnerabilities_tables = []            
            sorted_asc_vulns = vulns_list      
            sorted_asc_vulns.sort(key=lambda x: float(x["<<vulnerability_risk_score>>"]),reverse = True)
            for vunl in sorted_asc_vulns:
                # Upper Case letters
                vunl['<<vulnerability_name_upper>>'] = vunl['<<vulnerability_name>>'].upper()
            visible_mode_win32com = True
            f_name = file_name_no_ext+".xlsx"
            filename = os.path.join(output_folder,f_name) 
            logging.info(filename)
            sheetname = 'Vulnerabilidades'
            excel = win32.gencache.EnsureDispatch('Excel.Application')
            excel.Visible = visible_mode_win32com
            excel.DisplayAlerts = False
            wb = excel.Workbooks.Add()
            wb.Worksheets(1).Name = sheetname
            wb.Worksheets(1).Range("A1").FormulaR1C1 = "VULNERABILIDAD"
            wb.Worksheets(1).Range("B1").FormulaR1C1 = "NIVEL DE RIESGO"
            wb.Worksheets(1).Range("C1").FormulaR1C1 = "ESTADO"
            wb.Worksheets(1).Range("D1").FormulaR1C1 = "OBSERVACIONES"
            wb.Worksheets(1).ListObjects.Add(win32.constants.xlSrcRange, wb.Worksheets(1).Range("$A$1:$D$2"), win32.constants.xlYes,win32.constants.xlYes).Name = "TABLAVULNS"
            wb.Worksheets(1).ListObjects("TABLAVULNS").TableStyle = "TableStyleLight8"
            format_range_remediation_status("C2",wb)
            format_range_risk_leve("B2",wb)
            # Fill format
            index = 2
            for vunl in sorted_asc_vulns:
               wb.Worksheets(1).Range("A"+str(index)).Value = vunl['<<vulnerability_name_upper>>'] 
               index = index + 1 
            # Set risk levels
            index = 2
            for vunl in sorted_asc_vulns:
               wb.Worksheets(1).Range("B"+str(index)).Value =  get_risk_level_from_score(vunl['<<vulnerability_risk_score>>']) 
               index = index + 1 
            # Set status default
            index = 2
            for vunl in sorted_asc_vulns:
               wb.Worksheets(1).Range("C"+str(index)).Value =  "NO REMEDIADA"
               index = index + 1 
            # Adjust table wdith cells to fit

            wb.Worksheets(1).ListObjects("TABLAVULNS").Range.Columns.AutoFit()
            wb.SaveAs(filename)    
            wb.Close(False)
            excel.Application.Quit() 
        except Exception as e: 
            print(e)
            logging.info(e) 


def rgbToInt(rgb):
    colorInt = rgb[0] + (rgb[1] * 256) + (rgb[2] * 256 * 256)
    return colorInt

def covx_to_pdf(infile, outfile, word):
    wdFormatPDF = 17
    doc = word.Documents.Open(infile)
    doc.SaveAs(outfile, FileFormat=wdFormatPDF)
    doc.Close(False)

def printx_to_pdf(infile, outfile, word):
    wdFormatPDF = 17
    doc = word.Documents.Open(infile)
    newPrinter = "Microsoft Print to PDF"
    oldPrinter = word.ActivePrinter
    ActivePrinter = newPrinter
    doc.PrintOut(OutputFileName=outfile)
    word.ActivePrinter = oldPrinter
    doc.Close(False)

def count_inrange(list1, l, r):     
    # x for x in list1 is same as traversal in the list
    # the if condition checks for the number of numbers in the range
    # l to r
    # the return is stored in a list
    # whose length is the answer
    return len(list(x for x in list1 if l <= x <= r))
 

def merge_docx1(files, final_docx_name, visible_mode_win32com, output_folder):
    # Start word application
    word = win32.gencache.EnsureDispatch("Word.Application")
    word.Visible = visible_mode_win32com
    
    word.DisplayAlerts = False
    # New blank document
    new_document = word.Documents.Add()
    for fn in files:
        # Open each file to be merged, copy the contents to the clipboard, and then close the file
        fn = abspath(fn)
        temp_document = word.Documents.Open(fn)
        word.Selection.WholeStory()
        word.Selection.Copy()
        temp_document.Close(False)
        # Paste to the end of the new document
        new_document.Range()
        word.Selection.Delete()
        word.Selection.Paste() 
    clip.OpenClipboard()
    clip.EmptyClipboard()
    clip.CloseClipboard()
    # Save the final file and close the Word application
    new_document.SaveAs(os.path.join(output_folder,final_docx_name))    
    new_document.Close(False)
    doc = docx.Document(os.path.join(output_folder,final_docx_name))
    
    for table in doc.tables:
        try:
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="4C4C4C"/>'.format(nsdecls('w')))
            table.cell(0, 0)._tc.get_or_add_tcPr().append(shading_elm_1)
            shading_elm_2 = parse_xml(r'<w:shd {} w:fill="717171"/>'.format(nsdecls('w')))
            table.cell(0, 1)._tc.get_or_add_tcPr().append(shading_elm_2)
            if table.cell(0, 2).text != "-":
               risk_score_table = float(table.cell(0, 2).text)
            else:
               risk_score_table = 0    
            paragraph = table.cell(0, 2).paragraphs[0]
            paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            if(risk_score_table >= 0.1 and  risk_score_table <= 3.9):
                shading_elm_3 = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
                table.cell(0, 2)._tc.get_or_add_tcPr().append(shading_elm_3)
            elif(risk_score_table >= 4.0 and  risk_score_table <= 6.9):
                shading_elm_3 = parse_xml(r'<w:shd {} w:fill="FFC000"/>'.format(nsdecls('w')))
                table.cell(0, 2)._tc.get_or_add_tcPr().append(shading_elm_3)
            elif(risk_score_table >= 7.0 and  risk_score_table <= 8.9):
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255,255,255)
                shading_elm_3 = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w')))
                table.cell(0, 2)._tc.get_or_add_tcPr().append(shading_elm_3)
            elif(risk_score_table >= 9.0 and  risk_score_table <= 10): 
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255,255,255)           
                shading_elm_3 = parse_xml(r'<w:shd {} w:fill="C00000"/>'.format(nsdecls('w')))
                table.cell(0, 2)._tc.get_or_add_tcPr().append(shading_elm_3)
            elif(risk_score_table <= 0): 
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255,255,255)  
                shading_elm_3 = parse_xml(r'<w:shd {} w:fill="ADADAD"/>'.format(nsdecls('w')))
                table.cell(0, 2)._tc.get_or_add_tcPr().append(shading_elm_3)
            shading_elm_4 = parse_xml(r'<w:shd {} w:fill="717171"/>'.format(nsdecls('w')))
            table.cell(5, 0)._tc.get_or_add_tcPr().append(shading_elm_4)
            shading_elm_5 = parse_xml(r'<w:shd {} w:fill="717171"/>'.format(nsdecls('w')))
            table.cell(6, 0)._tc.get_or_add_tcPr().append(shading_elm_5)
            shading_elm_6 = parse_xml(r'<w:shd {} w:fill="717171"/>'.format(nsdecls('w')))
            table.cell(7, 0)._tc.get_or_add_tcPr().append(shading_elm_6)
            shading_elm_7 = parse_xml(r'<w:shd {} w:fill="717171"/>'.format(nsdecls('w')))
            table.cell(8, 0)._tc.get_or_add_tcPr().append(shading_elm_7)
            doc.save(os.path.join(output_folder,final_docx_name))
        except:
            doc.save(os.path.join(output_folder,final_docx_name))    

def split_text_before_point(text):

  parts = text.split('.')

  if len(parts) <= 1:
    return ''
  if len(parts) > 1 and parts[1] == '':
    return parts[0]+'.'
  elif len(parts) > 1 and parts[1][0] == ')':
    return parts[0]+').'
  else:
    return parts[0]+'.'

def sow_generation(wordapp, data, sow_targets_ips_string, sow_targets_urls, tmp_directory):
  

  dn = os.path.dirname(os.path.abspath(sys.argv[0]))

  # SOW GENERATION
   
  if data["<<analysis_type>>"] == "DYNAMIC VULNERABILITIES ANALYSIS OF DESKTOP APPLICATION":
    no_targets = data["<<no_targets>>"]
    concordancia_1 =  'a los sistemas' if int(no_targets) > 1 else 'al sistema'
    concordancia_2 =  'de las aplicaciones' if int(no_targets) > 1 else 'de la aplicación'
    concordancia_3 =  'a las aplicaciones' if int(no_targets) > 1 else 'a la aplicación'
    Dict = dict({'<<Nombre_del_aplicativo_portada>>': str(data['<<name_app>>'] + ' - ' + data['<<analysis_version_format_01>>']),
   '<<Fecha_mes_y_año>>':data['<<date_format_02>>'], 
   '<<Folio>>':data['<<analysis_id>>'],
   '<<Fecha_ddmmaa_encabezado>>':data['<<request_date_format_02>>'],
   '<<request_folio>>':data['<<request_folio>>'],
   '<<Folio>>':data['<<analysis_id>>'],
   '<<analysis_version_format_02>>': data['<<analysis_version_format_02>>'],
   '<<Concordancia_1>>':concordancia_1,
   '<<Nombre_del_aplicativo_En_antecedentes>>':data['<<name_app>>'],
   '<<Nombre_del_servidor>>':data['<<app_component>>'].replace("http://", "").replace("https://", ""),
   '<<Nombre_del_aplicativo_Tabla>>':data['<<name_app>>'], 
   '<<Fechas_de_inicio>>': data['<<start_date>>'],
   '<<Fecha_Fin>>': data['<<finish_date>>'],
   '<<Fecha_tentativa_de_inicio>>': data['<<start_date_planned>>'],
   '<<Fecha_límite_para_la_actividad>>': data['<<due_date>>'], 
   '<<Concordancia_2>>': concordancia_2, 
   '<<URL_Acuerdos_tabla3>>': sow_targets_urls,
   '<<supervisor>>': data['<<supervisor>>'],
   '<<supervisor_charge>>': data['<<supervisor_charge>>'],
   '<<supervisor_charge>>': data['<<supervisor_charge>>'],
   '<<client_responsible>>': data['<<client_responsible>>'],
   '<<client_responsible_charge>>': data['<<client_responsible_charge>>'],
   '<<client_company>>': data['<<client_company>>'],
   '<<company>>': data['<<company>>'],   
   '<<Concordancia_3>>':concordancia_3})
  else:
    no_targets = data["<<no_targets>>"]
    concordancia_1 =  'a los sistemas' if int(no_targets) > 1 else 'al sistema'
    concordancia_2 =  'de los portales' if int(no_targets) > 1 else 'del portal'
    concordancia_3 =  'a los portales' if int(no_targets) > 1 else 'al portal'
    Dict = dict({'<<Nombre_del_aplicativo_portada>>': str(data['<<name_app>>'] + ' - ' + data['<<analysis_version_format_01>>']),
   '<<Fecha_mes_y_año>>':data['<<date_format_02>>'], 
   '<<Folio>>':data['<<analysis_id>>'],
   '<<Fecha_ddmmaa_encabezado>>':data['<<request_date_format_02>>'],
   '<<Dirección_IP>>':sow_targets_ips_string,
   '<<request_folio>>':data['<<request_folio>>'],
   '<<Folio>>':data['<<analysis_id>>'],
   '<<analysis_version_format_02>>': data['<<analysis_version_format_02>>'],
   '<<Concordancia_1>>':concordancia_1,
   '<<Nombre_del_aplicativo_En_antecedentes>>':data['<<name_app>>'],
   '<<Nombre_del_servidor>>':data['<<app_component>>'].replace("http://", "").replace("https://", ""),
   '<<Nombre_del_aplicativo_Tabla>>':data['<<name_app>>'], 
   '<<Fechas_de_inicio>>': data['<<start_date>>'],
   '<<Fecha_Fin>>': data['<<finish_date>>'],
   '<<Fecha_tentativa_de_inicio>>': data['<<start_date_planned>>'],
   '<<Fecha_límite_para_la_actividad>>': data['<<due_date>>'], 
   '<<Concordancia_2>>': concordancia_2, 
   '<<URL_Acuerdos_tabla3>>': sow_targets_urls,
   '<<supervisor>>': data['<<supervisor>>'],
   '<<supervisor_charge>>': data['<<supervisor_charge>>'],
   '<<supervisor_charge>>': data['<<supervisor_charge>>'],
   '<<client_responsible>>': data['<<client_responsible>>'],
   '<<client_responsible_charge>>': data['<<client_responsible_charge>>'],
   '<<client_company>>': data['<<client_company>>'],
   '<<company>>': data['<<company>>'],   
   '<<Concordancia_3>>':concordancia_3})
  
  sow_template = os.path.join(dn,'templates',data['<<sow_report_template_filename>>'])
  sow_file_name = 'SOW - {}-{} {}'.format(data['<<analysis_id>>'],data['<<name_app>>'],data['<<analysis_version_format_01>>'])
  sow_full_file_name = os.path.join(dn,tmp_directory,sow_file_name+'.docx')
  doc = wordapp.Documents.Open(sow_template)
  doc.Activate()

  wordapp.Selection.GoTo(win32.constants.wdGoToPage, win32.constants.wdGoToAbsolute, "2")

  for From in Dict.keys():
      wordapp.ActiveWindow.ActivePane.View.SeekView =win32.constants.wdSeekMainDocument
      wordapp.Selection.Find.Execute(From, False, False, False, False, False, True, win32.constants.wdFindContinue, False, Dict[From], win32.constants.wdReplaceAll) 
      wordapp.ActiveWindow.ActivePane.View.SeekView = win32.constants.wdSeekCurrentPageHeader
      wordapp.Selection.Find.Execute(From, False, False, False, False, False, True, win32.constants.wdFindContinue, False, Dict[From], win32.constants.wdReplaceAll)     
  wordapp.ActiveWindow.ActivePane.View.SeekView =win32.constants.wdSeekMainDocument

  # Remove format hipertext in character space
  wordapp.Selection.Find.ClearFormatting
  wordapp.Selection.Find.Replacement.ClearFormatting
  wordapp.Selection.Find.Replacement.Font.Bold = False
  wordapp.Selection.Find.Replacement.Font.Italic = False
  wordapp.Selection.Find.Replacement.Font.Underline = win32.constants.wdUnderlineNone
  wordapp.Selection.Find.Replacement.Font.Color = win32.constants.wdColorAutomatic
  wordapp.Selection.Find.Text = "<space>"
  wordapp.Selection.Find.Replacement.Text = " y "
  wordapp.Selection.Find.Forward = True
  wordapp.Selection.Find.Wrap = win32.constants.wdFindContinue
  wordapp.Selection.Find.Format = True
  wordapp.Selection.Find.MatchCase = False
  wordapp.Selection.Find.MatchWholeWord = False
  wordapp.Selection.Find.MatchWildcards = False
  wordapp.Selection.Find.MatchSoundsLike = False
  wordapp.Selection.Find.MatchAllWordForms = False
  wordapp.Selection.Find.Execute(Replace=win32.constants.wdReplaceAll)

  wordapp.Selection.Find.ClearFormatting
  wordapp.Selection.Find.Replacement.ClearFormatting
  wordapp.Selection.Find.Replacement.Font.Bold = False
  wordapp.Selection.Find.Replacement.Font.Italic = False
  wordapp.Selection.Find.Replacement.Font.Underline = win32.constants.wdUnderlineNone
  wordapp.Selection.Find.Replacement.Font.Color = win32.constants.wdColorAutomatic
  wordapp.Selection.Find.Text = "<coma>"
  wordapp.Selection.Find.Replacement.Text = ", "
  wordapp.Selection.Find.Forward = True
  wordapp.Selection.Find.Wrap = win32.constants.wdFindContinue
  wordapp.Selection.Find.Format = True
  wordapp.Selection.Find.MatchCase = False
  wordapp.Selection.Find.MatchWholeWord = False
  wordapp.Selection.Find.MatchWildcards = False
  wordapp.Selection.Find.MatchSoundsLike = False
  wordapp.Selection.Find.MatchAllWordForms = False
  wordapp.Selection.Find.Execute(Replace=win32.constants.wdReplaceAll)

  wordapp.Selection.GoTo(win32.constants.wdGoToPage, win32.constants.wdGoToAbsolute, "1")
  wordapp.ActiveDocument.SaveAs(sow_full_file_name)
  doc.Close(False)

def add_bad_practices(wordapp, bad_practices):
  if bad_practices:
     bad_practices_list = []
     try:
        for item in bad_practices:                           
           bad_practices_list.append(item)         
        wordapp.Selection.HomeKey(Unit=win32.constants.wdStory)
        wordapp.Selection.Find.Execute('<<bad_practices_list>>')
        wordapp.Selection.Range.ListFormat.ApplyListTemplateWithLevel(ListTemplate = wordapp.ListGalleries(win32.constants.wdBulletGallery).ListTemplates(1), ContinuePreviousList= True, ApplyTo = win32.constants.wdListApplyToWholeList, DefaultListBehavior= win32.constants.wdWord10ListBehavior)
        wordapp.Selection.Font.Name = "Montserrat"
        wordapp.Selection.Text = '\n'.join(bad_practices_list)
     except Exception as e: 
             print(e)
  else: 
      wordapp.Selection.GoTo(What=win32.constants.wdGoToBookmark, Name = "BadBookMark").Delete() 

def add_qa_vulnerabilities(wordapp, qa_data):
  if qa_data:
    qa_vuln_list = []
    try:
        for item in qa_data:
            qa_vuln_list.append(item) 
        wordapp.Selection.HomeKey(Unit=win32.constants.wdStory)
        wordapp.Selection.Find.Execute('<<qa_vulnerabilities_list>>') 
        wordapp.Selection.Range.ListFormat.ApplyListTemplateWithLevel(ListTemplate = wordapp.ListGalleries(win32.constants.wdBulletGallery).ListTemplates(1), ContinuePreviousList= True, ApplyTo = win32.constants.wdListApplyToWholeList, DefaultListBehavior= win32.constants.wdWord10ListBehavior)
        wordapp.Selection.Font.Name = "Montserrat"
        wordapp.Selection.Text = '\n'.join(qa_vuln_list)
    except Exception as e: 
        print(e) 
  else:
    wordapp.Selection.GoTo(What=win32.constants.wdGoToBookmark, Name = "QABookMark").Delete() 

def generate_vulns_tablefile(wordapp, visible_mode_win32com, sorted_asc_vulns, full_vulns_table_file_name, tmp_directory_path, template_path):
    count = 1
    vulnerabilities_tables = []   
    for vunl in sorted_asc_vulns:                     
        wordapp = win32.gencache.EnsureDispatch("Word.Application")
        wordapp.Visible = visible_mode_win32com
        wordapp.DisplayAlerts = False
        doc = wordapp.Documents.Open(template_path)
        doc.Activate()
        for From in vunl.keys():
            wordapp.Selection.HomeKey(Unit=win32.constants.wdStory) 
            item_replace = str(vunl[From])
            if From == "<<vulnerability_ports>>":
                try:
                    item_replace = str(int(vunl[From]))
                except ValueError:
                    print("That's not an int!")
            elif not vunl[From]: # None
                item_replace = "-"
            elif vunl[From] == 0:
                item_replace = "-"  
            if From != "<<vulnerability_evidence_note>>" and From !="<<vulnerability_evidence_image_path>>" and From !="<<vulnerability_evidences>>":
                try:
                    wordapp.Selection.Find.Execute(From) 
                    wordapp.Selection.Text = item_replace
                except Exception as e: 
                    print(e)
            wordapp.Selection.EndKey(Unit=win32.constants.wdStory)
        try:  
            for evidence in vunl['<<vulnerability_evidences>>'] :  
                wordapp.Selection.EndKey(Unit=win32.constants.wdStory)       
                wordapp.Selection.TypeText(Text='\r') # New paragraph insted of Typeparagrah
                image_path = evidence['<<vulnerability_evidence_image_path>>'] 
                wordapp.Selection.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphCenter                
                shape = wordapp.Selection.InlineShapes.AddPicture(FileName=image_path,LinkToFile=False, SaveWithDocument=True )
                shape.LockAspectRatio = True
                #wordapp.Selection.EndOf
                wordapp.Selection.Collapse(Direction=win32.constants.wdCollapseEnd)
                wordapp.Selection.TypeText(Text='\r') # New paragraph insted of Typeparagrah
                wordapp.Selection.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphJustify
                #shape.Width = 450; # Change width works 
                wordapp.Selection.TypeText(Text=evidence['<<vulnerability_evidence_note>>'])
        except Exception as e: 
            print(e)
        if (count  <  len(sorted_asc_vulns)):
            wordapp.Selection.EndKey(Unit=win32.constants.wdStory) 
            wordapp.Selection.InsertBreak(Type=win32.constants.wdPageBreak) 
        
        file_name = os.path.join(tmp_directory_path,'Table_'+ str(count)+ ".docx")
        vulnerabilities_tables.append(file_name)
        wordapp.ActiveDocument.SaveAs(file_name)
        doc.Close(SaveChanges=True)
        count = count + 1
    return vulnerabilities_tables
    
def generate_report(data,visible_mode_win32com,tmp_directory, outputs_directory):
   if not os.path.exists(tmp_directory):
    os.makedirs(tmp_directory)

   if '_id' in data:
      del data['_id']      
   if data["<<vulnerabilities>>"]:
      for i in range(len(data["<<vulnerabilities>>"])):       
          for k in range(len(data["<<vulnerabilities>>"][i]["<<vulnerability_evidences>>"])):
              vulnerability_name = unidecode.unidecode(data["<<vulnerabilities>>"][i]["<<vulnerability_name>>"]).lower().replace(" ", "_").replace("/", "_")               
              image = os.path.join(tmp_directory,"vuln_{}_{}_ev_{}.png".format(str(i+1),vulnerability_name,str(k+1)))           
              fh = open(image, "wb")
              fh.write(base64.b64decode(data["<<vulnerabilities>>"][i]["<<vulnerability_evidences>>"][k]["<<vulnerability_evidence_image_path>>"]))
              fh.close()
              data["<<vulnerabilities>>"][i]["<<vulnerability_evidences>>"][k]["<<vulnerability_evidence_image_path>>"] = image            
   
   ## Transforms array lists to one string with /r characters
   for vuln in data["<<vulnerabilities>>"]:
       vuln["<<vulnerability_references_string>>"] = '\r'.join(vuln["<<vulnerability_references>>"])
       vuln.pop("<<vulnerability_references>>", None)
       vuln["<<vulnerability_paths_string>>"] = '\r'.join(vuln["<<vulnerability_paths>>"])
       vuln.pop("<<vulnerability_paths>>", None)
   # "<<analysis_revision_01>>"
   data["<<analysis_revision_01>>"] = "{:.1f}".format(data["<<analysis_revision_number>>"])
   # Date formats manipulation
   start_date_planned = datetime.strptime(data["<<start_date_planned>>"], '%d/%m/%Y')
   data["<<start_date_planned>>"] = str(start_date_planned.strftime("%d de %B de %Y"))   
   due_date = datetime.strptime(data["<<due_date>>"], '%d/%m/%Y')
   data["<<due_date>>"] = str(due_date.strftime("%d de %B de %Y"))   
   start_date = datetime.strptime(data["<<start_date>>"], '%d/%m/%Y')
   data["<<start_date>>"] = str(start_date.strftime("%d de %B de %Y"))   
   request_date = datetime.strptime(data["<<request_date>>"], '%d/%m/%Y')
   data["<<request_date>>"] = str(request_date.strftime("%d de %B de %Y"))   
   finish_date = datetime.strptime(data["<<finish_date>>"], '%d/%m/%Y')
   data["<<finish_date>>"] = str(finish_date.strftime("%d de %B de %Y"))
   generation_date = datetime.strptime(data["<<generation_date>>"], '%d/%m/%Y')
   data["<<date_format_02>>"] = str(generation_date.strftime("%m|%Y"))
   data["<<request_date_format_01>>"] = str(request_date.strftime("%d de %B de %Y"))
   data["<<request_date_format_02>>"] = str(request_date.strftime("%d/%m/%Y"))
   
   # No Targets
   data["<<no_targets>>"] = len(data["<<scope>>"])
   # Add calculated fields 
   data['<<analysis_version_format_01>>'] = "{}{} {}".format(str(data['<<analysis_version>>']),num2ordinalabrev(int(data['<<analysis_version>>'])) ,"análisis")
   data['<<analysis_version_format_02>>'] = "{}".format(num2ordinalapocade(int(data['<<analysis_version>>'])))


   # Opening JSON file
   #with open(analysis_filename, encoding='utf-8') as json_file:
   #    data = json.loads(json.load(json_file))
   
   #template_file_path = os.path.join(dn,'templates',data['<<main_report_template_filename>>'])
   dn = os.path.dirname(os.path.abspath(sys.argv[0]))
   
   
   

   template_file_path = os.path.join(dn,'templates',data['<<main_report_template_filename>>']).replace('\r', '')   
   name_file = data['<<analysis_id>>'] + ' ' + data ['<<name_app>>'] + ' - ' + data['<<analysis_version_format_01>>'] + ".docx"
   base_name_file = data['<<analysis_id>>'] + ' ' + data ['<<name_app>>'] + ' - ' + data['<<analysis_version_format_01>>']
   name_file = name_file.replace("/", "-").replace('\r', '')
   
   wordapp = win32.gencache.EnsureDispatch("Word.Application",pythoncom.CoInitialize())

   wordapp.Visible = visible_mode_win32com
   wordapp.DisplayAlerts = False
   
   
   doc = wordapp.Documents.Open(template_file_path)
   doc.Activate()
   wordapp.Selection.HomeKey(Unit=win32.constants.wdStory)

   sow_target_ip_list = []
   sow_target_url_list = []
   sow_targets_ips_string = []
   sow_targets_urls = []


   try:
      wordapp.Selection.Find.Execute("<<scope_table>>") 
      if data['<<scope>>']:
        for target in data['<<scope>>']: 
            try:   
               
               if data["<<analysis_type>>"] == "DYNAMIC VULNERABILITIES ANALYSIS OF DESKTOP APPLICATION":
                 # Create table structure
                 wordapp.ActiveDocument.Tables.Add(Range=wordapp.Selection.Range, NumRows=2, NumColumns= 2, DefaultTableBehavior=win32.constants.wdWord9TableBehavior, AutoFitBehavior= win32.constants.wdAutoFitFixed)
                 wordapp.Selection.Tables(1).PreferredWidthType = win32.constants.wdPreferredWidthPoints
                 wordapp.Selection.Tables(1).PreferredWidth = 368.503937008 # 13 Centimeters to Point
                 wordapp.Selection.Tables(1).Columns(1).SetWidth(ColumnWidth=141.5, RulerStyle= win32.constants.wdAdjustNone)
                 wordapp.Selection.Tables(1).Rows.Alignment = win32.constants.wdAlignRowCenter   
                 # Format cells
                 wordapp.Selection.Tables(1).Cell(1, 1).Range.Shading.BackgroundPatternColor = -570376193
                 wordapp.Selection.Tables(1).Cell(1, 1).Range.Font.Bold = win32.constants.wdToggle
                 wordapp.Selection.Tables(1).Cell(1, 1).Range.Font.Color = rgbToInt((255,255,255)) 
                 wordapp.Selection.Tables(1).Cell(2, 1).Range.Shading.BackgroundPatternColor = -570376193
                 wordapp.Selection.Tables(1).Cell(2, 1).Range.Font.Bold = win32.constants.wdToggle
                 wordapp.Selection.Tables(1).Cell(2, 1).Range.Font.Color = rgbToInt((255,255,255)) 
                 # Populate cells
                 wordapp.Selection.Tables(1).Cell(1, 1).Range.Text = "Descripción"
                 wordapp.Selection.Tables(1).Cell(2, 1).Range.Text = "Sistema Operativo"
                 wordapp.Selection.Tables(1).Cell(1, 2).Range.Text = target["<<target_description>>"]
                 wordapp.Selection.Tables(1).Cell(2, 2).Range.Text = target["<<target_operative_system>>"]
                 sow_target_url_list.append(data["<<app_component>>"])
               else:
                 # Create table structure
                 wordapp.ActiveDocument.Tables.Add(Range=wordapp.Selection.Range, NumRows=4, NumColumns= 2, DefaultTableBehavior=win32.constants.wdWord9TableBehavior, AutoFitBehavior= win32.constants.wdAutoFitFixed)
                 wordapp.Selection.Tables(1).PreferredWidthType = win32.constants.wdPreferredWidthPoints
                 wordapp.Selection.Tables(1).PreferredWidth = 368.503937008 # 13 Centimeters to Point
                 wordapp.Selection.Tables(1).Columns(1).SetWidth(ColumnWidth=141.5, RulerStyle= win32.constants.wdAdjustNone)
                 wordapp.Selection.Tables(1).Rows.Alignment = win32.constants.wdAlignRowCenter   
                 # Format cells
                 wordapp.Selection.Tables(1).Cell(1, 1).Range.Shading.BackgroundPatternColor = -570376193
                 wordapp.Selection.Tables(1).Cell(1, 1).Range.Font.Bold = win32.constants.wdToggle
                 wordapp.Selection.Tables(1).Cell(1, 1).Range.Font.Color = rgbToInt((255,255,255)) 
                 wordapp.Selection.Tables(1).Cell(2, 1).Range.Shading.BackgroundPatternColor = -570376193
                 wordapp.Selection.Tables(1).Cell(2, 1).Range.Font.Bold = win32.constants.wdToggle
                 wordapp.Selection.Tables(1).Cell(2, 1).Range.Font.Color = rgbToInt((255,255,255)) 
                 wordapp.Selection.Tables(1).Cell(3, 1).Range.Shading.BackgroundPatternColor = -570376193
                 wordapp.Selection.Tables(1).Cell(3, 1).Range.Font.Bold = win32.constants.wdToggle
                 wordapp.Selection.Tables(1).Cell(3, 1).Range.Font.Color = rgbToInt((255,255,255)) 
                 wordapp.Selection.Tables(1).Cell(4, 1).Range.Shading.BackgroundPatternColor = -570376193
                 wordapp.Selection.Tables(1).Cell(4, 1).Range.Font.Bold = win32.constants.wdToggle
                 wordapp.Selection.Tables(1).Cell(4, 1).Range.Font.Color =  rgbToInt((255,255,255)) 
                 # Populate cells
                 wordapp.Selection.Tables(1).Cell(1, 1).Range.Text = "Dirección IP"
                 wordapp.Selection.Tables(1).Cell(2, 1).Range.Text = "URL"
                 wordapp.Selection.Tables(1).Cell(3, 1).Range.Text = "Descripción"
                 wordapp.Selection.Tables(1).Cell(4, 1).Range.Text = "Sistema Operativo"
                 wordapp.Selection.Tables(1).Cell(1, 2).Range.Text = target["<<target_ip>>"]
                 wordapp.Selection.Tables(1).Cell(2, 2).Range.Text = target["<<target_url>>"]
                 wordapp.Selection.Tables(1).Cell(3, 2).Range.Text = target["<<target_description>>"]
                 wordapp.Selection.Tables(1).Cell(4, 2).Range.Text = target["<<target_operative_system>>"]
                 sow_target_url_list.append(target["<<target_url>>"])
                 sow_target_ip_list.append(target["<<target_ip>>"])                       
               wordapp.Selection.MoveDown(Unit=win32.constants.wdParagraph, Count=12)
               wordapp.Selection.TypeText(Text="\r\n")
            except Exception as e:
               print(e) 
   except Exception as e:
      print(e) 

   
   wordapp.Selection.HomeKey(Unit=win32.constants.wdStory)
   
   sow_targets_ips_dict = list(dict.fromkeys(sow_target_ip_list))
   sow_targets_urls_dict = list(dict.fromkeys(sow_target_url_list))
   if len(sow_targets_ips_dict) > 1:      
      sow_targets_ips_string  = ", ".join(sow_targets_ips_dict[:-1]) +" y "+sow_targets_ips_dict[-1]
   elif len(sow_targets_ips_dict) == 1:
      sow_targets_ips_string = sow_target_ip_list[0]
   else:
      sow_targets_ips_string = ""

   if len(sow_targets_urls_dict) > 1:      
      sow_targets_urls  = "<coma>".join(sow_targets_urls_dict[:-1]) +"<space>"+sow_targets_urls_dict[-1]
   elif len(sow_targets_urls_dict) == 1:
      sow_targets_urls = sow_target_url_list[0]
   else:
      sow_targets_urls = ""
   

   data['<<previous_analysis_version_format_03>>'] = ''
   try:
      data['<<previous_analysis_version_format_03>>'] = num2words(int(data['<<analysis_version>>'])-1, to='ordinal',lang='es ')
   except NotImplementedError:
      data['<<previous_analysis_version_format_03>>']  = num2words(1, lang='es ', to='ordinal')
   
   risk_scores = []
   risk_list = []
   risk_resume_list = []
   remediation_list = []
   impact_list = []
   for vunl in data['<<vulnerabilities>>']:               
       risk_scores.append(float(vunl["<<vulnerability_risk_score>>"]))              
       risk_list.append(split_text_before_point(vunl["<<vulnerability_risk>>"]))
       if vunl["<<vulnerability_clasification>>"]:
         impact_list.extend(vunl["<<vulnerability_clasification>>"].lower().replace(' y',',').replace('.','').split(", "))
       remediation_list.append(split_text_before_point(vunl["<<vulnerability_remediation>>"])) 
   
   if not data['<<vulnerabilities>>']:
      # The section 1 of document correspond cover page
      # section 2 corresponde text signs si le das click en encabezado veras numero seccio
      # Ver marcadores ocultos para ver los de los titulos      
      executive_resume = "Con el propósito de identificar las vulnerabilidades potenciales que pudieran comprometer la seguridad de la Información, se realizó el <<analysis_version_format_02>> análisis dinámico del aplicativo, sin embargo, no se identificaron nuevas vulnerabilidades."
      wordapp.Selection.Find.Execute('<<executive_resume>>') 
      wordapp.Selection.TypeText(Text=executive_resume)     
      wordapp.Selection.GoTo(What=win32.constants.wdGoToBookmark, Name = "ReportStructureBookMark")
      wordapp.Selection.Style = wordapp.ActiveDocument.Styles("Normal")
      if data['<<previous_analysis_version_format_03>>']:
        wordapp.Selection.TypeText(Text="Se observa que las vulnerabilidades identificadas en el <<previous_analysis_version_format_03>> análisis dinámico aplicado al portal web <<name_app>> se encuentran remediadas.\rPara dar continuidad en la seguridad del aplicativo, se recomiendan las siguientes acciones:\r")
      else: 
        wordapp.Selection.TypeText(Text="Se observa que las vulnerabilidades identificadas en el análisis dinámico aplicado al portal web <<name_app>> se encuentran remediadas.\rPara dar continuidad en la seguridad del aplicativo, se recomiendan las siguientes acciones:\r")
      list_suggestions_empty_report = ["Mantener actualizado los componentes del aplicativo, frameworks y API´s.","Actualizar periódicamente las credenciales de acceso y privilegios de usuarios que realizan cambios críticos en la infraestructura o manejan información sensible.","Realizar respaldos periódicos al aplicativo.","Mantener activo y respaldado los logs de cada componente del aplicativo."]
      wordapp.Selection.TypeParagraph
      wordapp.Selection.Font.Name = "Montserrat"
      wordapp.Selection.Range.ListFormat.ApplyListTemplateWithLevel(ListTemplate = wordapp.ListGalleries(win32.constants.wdBulletGallery).ListTemplates(1), ContinuePreviousList= True, ApplyTo = win32.constants.wdListApplyToWholeList, DefaultListBehavior= win32.constants.wdWord10ListBehavior)
      wordapp.Selection.ParagraphFormat.LeftIndent = 36
      wordapp.Selection.Text = '\n'.join(list_suggestions_empty_report)
      wordapp.Selection.Font.Name = "Montserrat"
      add_qa_vulnerabilities(wordapp,data['<<qa_vulnerabilities>>'])  
      add_bad_practices(wordapp,data['<<bad_practices_list>>'])  
      wordapp.Selection.GoTo(What=win32.constants.wdGoToBookmark, Name = "VulnesPartBookMark").Delete()  
      wordapp.Selection.GoTo(What=win32.constants.wdGoToBookmark, Name = "CriticityAndResumeBookMark").Delete()   
   else:
      if impact_list and not list(set(impact_list))[0] == "-":
          executive_resume = "Con el propósito de identificar las vulnerabilidades potenciales que pudieran comprometer la seguridad de la Información, se realizó el <<analysis_version_format_02>> análisis dinámico del aplicativo. Las vulnerabilidades identificadas podrían permitir a un atacante <<risk_resume_list>>. Estos se clasifican con un riesgo <<level_max>>, por lo que se considera que se deben realizar <<executive_resume_part1>> a la <<executive_resume_part2>> de los sistemas involucrados."
      else:
          executive_resume = "Con el propósito de identificar las vulnerabilidades potenciales que pudieran comprometer la seguridad de la Información, se realizó el <<analysis_version_format_02>> análisis dinámico del aplicativo. Las vulnerabilidades identificadas podrían permitir a un atacante <<risk_resume_list>>. Estos se clasifican con un riesgo <<level_max>>, por lo que se considera que se deben realizar <<executive_resume_part1>> a los sistemas involucrados."
      wordapp.Selection.Find.Execute('<<executive_resume>>') 
      wordapp.Selection.Text = executive_resume
   
   wordapp.ActiveDocument.TrackRevisions = False  # Maybe not need this (not really but why not)
   wordapp.Selection.GoTo(win32.constants.wdGoToPage, win32.constants.wdGoToAbsolute, "2")

   
   try:
      for From in data.keys():
          logging.info(From)
          if (From != '<<scope>>' and From != '<<qa_vulnerabilities>>' and From != '<<bad_practices_list>>' and From != '<<vulnerabilities>>'):
              wordapp.ActiveWindow.ActivePane.View.SeekView =win32.constants.wdSeekMainDocument
              wordapp.Selection.Find.Execute(From, False, False, False, False, False, True, win32.constants.wdFindContinue, False, data[From], win32.constants.wdReplaceAll) 
              wordapp.ActiveWindow.ActivePane.View.SeekView = win32.constants.wdSeekCurrentPageHeader              
              wordapp.Selection.Find.Execute(From, False, False, False, False, False, True, win32.constants.wdFindContinue, False, data[From], win32.constants.wdReplaceAll) 
          wordapp.ActiveWindow.ActivePane.View.SeekView =win32.constants.wdSeekMainDocument
   except Exception as e: 
           print(e)  
   

   
   wordapp.Selection.GoTo(win32.constants.wdGoToPage, win32.constants.wdGoToAbsolute, "2")


   

   logging.info("Debug point A")
   wordapp.ActiveDocument.SaveAs(os.path.join(dn,tmp_directory,name_file))
   doc.TablesOfContents(1).Update()
   wordapp.ActiveDocument.Save()
   doc.Close(SaveChanges=True)

   

   if data['<<vulnerabilities>>']:
        try:
            count = 1
            sorted_asc_vulns = data['<<vulnerabilities>>']
            sorted_asc_vulns.sort(key=lambda x: float(x["<<vulnerability_risk_score>>"]),reverse = True)

            for vunl in sorted_asc_vulns:
                # Upper Case letters
                vunl['<<vulnerability_name_upper>>'] = vunl['<<vulnerability_name>>'].upper()

            vulns_table_file_name =   'TB - {} {} {}{}'.format(data['<<analysis_id>>'],data ['<<name_app>>'],data['<<analysis_version_format_01>>'],".docx")
            vulns_table_file_name = vulns_table_file_name.replace("/", "-").replace('\r', '')
            full_vulns_table_file_name = os.path.join(dn,tmp_directory,vulns_table_file_name).replace('\r', '')
            tmp_directory_path = os.path.join(dn,tmp_directory)
            table_template = os.path.join(dn,'templates',data ['<<vuln_table_template_filename>>'])   
            vulnerabilities_tables = generate_vulns_tablefile(wordapp, visible_mode_win32com, sorted_asc_vulns, full_vulns_table_file_name, tmp_directory, table_template)
            merge_docx1(vulnerabilities_tables,vulns_table_file_name, visible_mode_win32com = visible_mode_win32com, output_folder = os.path.join(dn,tmp_directory))
            
            doca = wordapp.Documents.Open(os.path.join(dn,tmp_directory,name_file))
            doca.Activate()            
            wordapp.Selection.Find.Execute('<<vulnerabilities_tables>>') 
            wordapp.Selection.InsertFile(FileName=full_vulns_table_file_name, Range="", ConfirmConversions=False, Link=False, Attachment=False)
            # Delete terminal characters
            
            #wordapp.Selection.MoveLeft(Unit=win32.constants.wdCharacter, Count=2)  
            #wordapp.Selection.Delete(Unit=win32.constants.wdCharacter, Count=2)
            #wordapp.Selection.Find.Execute('<<break_page>>', False, False, False, False, False, True, win32.constants.wdFindContinue, False, "^m", win32.constants.wdReplaceAll) 
            
            #wordapp.Selection.InsertBreak(Type=win32.constants.wdSectionBreakContinuous)
            wordapp.Selection.MoveLeft(Unit=win32.constants.wdCharacter, Count=2) 
            wordapp.Selection.Delete(Unit=win32.constants.wdCharacter, Count=2)
            #wordapp.Selection.MoveRight(Unit=win32.constants.wdCharacter, Count=1)              
            #wordapp.Selection.MoveLeft(Unit=win32.constants.wdCharacter, Count=1) 
            #wordapp.Selection.Delete(Unit=win32.constants.wdCharacter, Count=3)
            #wordapp.Selection.TypeText(Text="^m")
            #wordapp.Selection.InsertBreak(Type=win32.constants.wdPageBreak)
            
            
            
            #wordapp.Selection.InsertBreak(Type=win32.constants.wdPageBreak)  
            # Go to start document
            wordapp.Selection.HomeKey(Unit=win32.constants.wdStory)
            wordapp.Selection.Find.Execute('<<level_max>>') 
            wordapp.Selection.Font.Bold = True             
            
            for item in risk_list:
              risk_resume_list.append(item.replace("Un atacante podría ", "").replace(".", ""))             

            wordapp.Selection.HomeKey(Unit=win32.constants.wdStory)
            wordapp.Selection.Find.Execute('<<risk_resume_list>>')
            logging.info("Debug point 1")            
            wordapp.Selection.Text = ', '.join(risk_resume_list)
            
            # Dictionary count
            dict_of_counts = {}
            dict_of_counts["INFORMATIVA"] = count_inrange(risk_scores,-1,0)
            dict_of_counts["BAJA"] = count_inrange(risk_scores,0.1,3.9)
            dict_of_counts["MEDIA"] = count_inrange(risk_scores,4.0,6.9)
            dict_of_counts["ALTA"] = count_inrange(risk_scores,7.0,8.9)
            dict_of_counts["CRÍTICA"] = count_inrange(risk_scores,9.0,10)

            max_score = max(risk_scores)
            max_level = 'BAJO'
            if (max_score >= 9.0):
                max_level = 'CRÍTICO'
            elif (max_score >= 7.0):
                max_level = 'ALTO'
            elif (max_score >= 4.0):
                max_level = 'MEDIO'
            elif (max_score >= 0.1):
                max_level = 'BAJO'
            elif (max_score <= 0.0):
                max_level = 'INFORMATIVO'

            impact_list = list(dict.fromkeys(impact_list))
            if impact_list:
                if len(impact_list) > 1:
                   if impact_list[-1].startswith('i'):
                    impact_string = ", ".join(impact_list[:-1]) +" e "+impact_list[-1]
                   else:
                    impact_string = ", ".join(impact_list[:-1]) +" y "+impact_list[-1]
                else: 
                    impact_string = impact_list[0]

            logging.info("Debung point 2")
            
            # Have to identify the index of the graph you want to handle
            if (wordapp.ActiveDocument.InlineShapes(1).Type == 12): # Is a chart
                chart_wb = wordapp.ActiveDocument.InlineShapes(1).Chart.ChartData.Workbook
                chart_wb.Application.Visible = visible_mode_win32com
                #wordapp.ActiveDocument.InlineShapes(1).Chart.ChartData.Activate()
                SourceSheet = chart_wb.ActiveSheet
                SourceSheet.Range("B2").Value2 = dict_of_counts["INFORMATIVA"]
                SourceSheet.Range("B3").Value2 = dict_of_counts["BAJA"]
                SourceSheet.Range("B4").Value2 = dict_of_counts["MEDIA"]
                SourceSheet.Range("B5").Value2 = dict_of_counts["ALTA"]
                SourceSheet.Range("B6").Value2 = dict_of_counts["CRÍTICA"]
            
                if (dict_of_counts["INFORMATIVA"] <= 0):
                    wordapp.ActiveDocument.InlineShapes(1).Chart.ChartGroups(1).FullCategoryCollection(1).IsFiltered = True
                if (dict_of_counts["BAJA"] <=  0):
                    wordapp.ActiveDocument.InlineShapes(1).Chart.ChartGroups(1).FullCategoryCollection(2).IsFiltered = True
                if (dict_of_counts["MEDIA"] <=  0):
                    wordapp.ActiveDocument.InlineShapes(1).Chart.ChartGroups(1).FullCategoryCollection(3).IsFiltered = True
                if (dict_of_counts["ALTA"] <=  0):
                    wordapp.ActiveDocument.InlineShapes(1).Chart.ChartGroups(1).FullCategoryCollection(4).IsFiltered = True
                if (dict_of_counts["CRÍTICA"] <=  0):
                    wordapp.ActiveDocument.InlineShapes(1).Chart.ChartGroups(1).FullCategoryCollection(5).IsFiltered = True
                
                # xlCellTypeBlanks =  4
                #SourceSheet.Range("A2:B5").SpecialCells(4).Delete()
            
                chart_wb.Close(True)
                wordapp.ActiveDocument.InlineShapes(1).Chart.Refresh
                #wordapp.ActiveDocument.InlineShapes(1).Chart.ChartGroups(1).FullCategoryCollection(1).IsFiltered = True
                #wordapp.ActiveDocument.InlineShapes(1).Chart.SeriesCollection(1).DataLabels.ShowValue = False
          
            wordapp.Selection.Find.Execute("<<level_max>>", False, False, False, False, False, True, win32.constants.wdFindContinue, False, max_level, win32.constants.wdReplaceAll) 
            if max_level == 'INFORMATIVO':
               text_part = 'acciones para mitigar las vulnerabilidades expuestas ya que estas podrían afectar en ciertas condiciones'
               wordapp.Selection.Find.Execute("<<executive_resume_part1>>", False, False, False, False, False, True, win32.constants.wdFindContinue, False, text_part, win32.constants.wdReplaceAll) 
            else:
               text_part = 'acciones inmediatas para mitigar las vulnerabilidades expuestas ya que estas podrían afectar directamente'
               wordapp.Selection.Find.Execute("<<executive_resume_part1>>", False, False, False, False, False, True, win32.constants.wdFindContinue, False, text_part, win32.constants.wdReplaceAll) 
            
            if impact_list:
                wordapp.Selection.Find.Execute("<<executive_resume_part2>>", False, False, False, False, False, True, win32.constants.wdFindContinue, False, impact_string, win32.constants.wdReplaceAll) 

            
            # Get the correct index of table
            #print(doca.Tables)            
            logging.info(doca.Tables(2).Cell(1, 1).Range.Text)
            logging.info(doca.Tables(2).Cell(3, 1).Range.Text)
            logging.info(doca.Tables(2).Rows.Count)
            
            doca.Tables(2).Rows.Add()
            doca.Tables(2).Cell(3, 1).Select() 
            wordapp.Selection.SelectRow() 
            wordapp.Selection.Cells.Delete(ShiftCells=win32.constants.wdDeleteCellsEntireRow)
            
            #doca.Tables(2).Cell(3, 2).Delete()
            #doca.Tables(2).Cell(3, 1).Delete()

            
            
            for vunl in sorted_asc_vulns:
                if(float(vunl['<<vulnerability_risk_score>>']) >= 0.1 and  float(vunl['<<vulnerability_risk_score>>']) <= 3.9):
                    vunl['<<level_risk>>'] = 'BAJA'
                    vunl['<<level_risk_olecolor>>'] = 65535
                    vunl['<<level_risk_text_olecolor>>'] = -16777216
                    vunl['<<level_risk_text_rgbcolor>>'] = RGBColor(255,255,255)
                elif(float(vunl['<<vulnerability_risk_score>>']) >= 4.0 and  float(vunl['<<vulnerability_risk_score>>']) <= 6.9):
                    vunl['<<level_risk>>'] = 'MEDIA'
                    vunl['<<level_risk_olecolor>>'] = 49407
                    vunl['<<level_risk_text_olecolor>>'] = -16777216
                    vunl['<<level_risk_text_rgbcolor>>'] = RGBColor(255,255,255)
                elif(float(vunl['<<vulnerability_risk_score>>']) >= 7.0 and  float(vunl['<<vulnerability_risk_score>>']) <= 8.9):
                    vunl['<<level_risk>>'] = 'ALTA'
                    vunl['<<level_risk_olecolor>>'] = 255
                    vunl['<<level_risk_text_olecolor>>'] = 16777215
                    vunl['<<level_risk_text_rgbcolor>>'] = RGBColor(0,0,0)
                elif(float(vunl['<<vulnerability_risk_score>>']) >= 9.0 and  float(vunl['<<vulnerability_risk_score>>']) <= 10): 
                    vunl['<<level_risk>>'] = 'CRÍTICA'
                    vunl['<<level_risk_olecolor>>'] = 192
                    vunl['<<level_risk_text_olecolor>>'] = 16777215
                    vunl['<<level_risk_text_rgbcolor>>'] =  RGBColor(0,0,0)
                elif(float(vunl['<<vulnerability_risk_score>>']) <= 0): 
                    vunl['<<level_risk>>'] = 'INFORMATIVA'
                    vunl['<<level_risk_olecolor>>'] = 11382189
                    vunl['<<level_risk_text_olecolor>>'] = -16777216
                    vunl['<<level_risk_text_rgbcolor>>'] = RGBColor(255,255,255)
                    
               
            index = 3
            for vunl in sorted_asc_vulns:
                doca.Tables(2).Cell(index, 1).Range.Text = vunl['<<level_risk>>']
                doca.Tables(2).Cell(index, 1).Shading.BackgroundPatternColor = vunl['<<level_risk_olecolor>>']
                doca.Tables(2).Cell(index, 1).Range.Font.Color = vunl['<<level_risk_text_olecolor>>']
                doca.Tables(2).Cell(index, 2).Range.Text =  "{0:3}. {1}".format(index-2,vunl['<<vulnerability_name>>']) 
                
                index = index + 1 
                doca.Tables(2).Rows.Add()
            doca.Tables(2).Cell(index, 1).Select() 
            wordapp.Selection.SelectRow() 
            wordapp.Selection.Cells.Delete(ShiftCells=win32.constants.wdDeleteCellsEntireRow)
            
            
            wordapp.Selection.HomeKey(Unit=win32.constants.wdStory)
            wordapp.Selection.Find.Execute('<<risk_list>>') 
            wordapp.Selection.Text = '\n'.join(risk_list)
            
            wordapp.Selection.HomeKey(Unit=win32.constants.wdStory)
            wordapp.Selection.Find.Execute('<<recomendation_list>>') 
            wordapp.Selection.Text = '\n'.join(remediation_list)
            
            add_qa_vulnerabilities(wordapp, data['<<qa_vulnerabilities>>'])
            add_bad_practices(wordapp, data['<<bad_practices_list>>'])
              
            doca.TablesOfContents(1).Update() 
            wordapp.ActiveDocument.Save()
            doca.Close(SaveChanges=True)
            
            sow_generation(wordapp, data, sow_targets_ips_string, sow_targets_urls, tmp_directory)

            # Add QA vulns to table vulns
            doct = os.path.join(dn,tmp_directory,full_vulns_table_file_name)
            add_qa_vulns_to_tablefile(doct,wordapp,data['<<qa_vulnerabilities>>'])
            wordapp.Quit()
           
            # Convert o PDF optional
            #try:
            #   covx_to_pdf(os.path.join(dn,tmp_directory,name_file), os.path.join(dn,tmp_directory,name_file).replace('.docx', '.pdf'),wordapp)
            #   covx_to_pdf(sow_full_file_name, sow_full_file_name.replace('.docx', '.pdf'),wordapp)
            #except Exception as e: 
            #   print(e)

            # Create Vuln Matrix
            matrix_file_name = 'MTZ - {}-{} {}'.format(data['<<analysis_id>>'],data['<<name_app>>'],data['<<analysis_version_format_01>>'])
            if data['<<vulnerabilities>>']:
              create_vuln_matrix(data['<<vulnerabilities>>'], tmp_directory, matrix_file_name) 
            
            
            for f in vulnerabilities_tables:
                # Delete temporal files generated
                if os.path.isfile(f): # this makes the code more robust
                        os.remove(f)
            
            if os.path.isfile(vulns_table_file_name): # this makes the code more robust
                        os.remove(vulns_table_file_name)

        except Exception as e: 
           print(e) 
   else:
    
    sow_generation(wordapp, data, sow_targets_ips_string, sow_targets_urls, tmp_directory)
    wordapp.Quit()

   try:
      #print("Deleting {}...".format(win32com.__gen_path__))
      shutil.rmtree(win32com.__gen_path__, ignore_errors=True) 
   except Exception as e: 
       print(e)

   
   # Remove file from zip temporary folder
   for item in os.listdir(tmp_directory):
      if item.endswith(".png"):
        os.remove(os.path.join(tmp_directory, item))

   zip_name = os.path.join(outputs_directory,base_name_file)
   zip_path = os.path.join(outputs_directory,"{}.zip".format(base_name_file))

   files_in_directory = os.listdir(tmp_directory)
   shutil.make_archive(zip_name, 'zip', tmp_directory)

   logging.info("Debug point")
   
   for files in os.listdir(tmp_directory):
       path = os.path.join(tmp_directory, files)
       try:
          shutil.rmtree(path)
       except OSError:
           os.remove(path)


def markdown_tojson(markdown_filename):
   
   # Create target directory & all intermediate directories if don't exists
   if not os.path.isfile(markdown_filename):
       print('No existe ese archivo.')
       exit()
   
   markdown_file  = open(markdown_filename,mode="r", encoding="utf-8") 
   markdown_file_data = markdown_file.read()
   markdown_file.close()
   html_start = '<html><body>'
   html_end = "</body></html>"
   html_str = mistune.markdown(markdown_file_data)
   complete_html = html_start+html_str+html_end
   
   Html_file= open(os.path.join(dn,"outputs",'temp.html') ,"w")
   Html_file.write(complete_html)
   Html_file.close()
   
   soup = BeautifulSoup(complete_html, "lxml")
   
   invalid_tags = ['b','a', 'i', 'u']
   
   for tag in invalid_tags: 
       for match in soup.findAll(tag):
           match.replaceWithChildren()
   
   with open(os.path.join(dn,"outputs",'temp-nohref.html'), "w") as file:
       file.write(str(soup))
   
   relevant_document_html_tags = ['h1','h2','h3','h4','h5','h6','p']
   vuln_tags = ['h2','h6','p']
   data_structured = {}
   
   
   for element in soup.find_all('h1'):  
       if  element.text == 'App information':  
           tag = element.find_next_sibling(relevant_document_html_tags)
           while tag.name == 'h6' or tag.name == 'p':
               content = "" 
               if tag.name == 'h6': 
                   key = tag.text                
               elif tag.name == 'p':
                   subtag = tag                
                   while subtag.name == 'p':
                       content = content + subtag.text     
                       subtag = subtag.find_next_sibling(relevant_document_html_tags)
               data_structured["<<"+key+">>"] = content
               tag = tag.find_next_sibling(relevant_document_html_tags)
   
   
   data_structured['<<vulnerabilities>>'] = []
   vuln = {}  
   key = ""
   content = ""
   for element in soup.find_all('h1',text ='Vulnerabilities'):
           subelement = element.find_next_sibling()  
           while subelement is not None:
               if subelement.name == 'h1': 
                   break # Stop list all elements final state
               elif subelement.name == 'h2':
                   if not vuln:
                       vuln["<<vulnerability_name>>"] = subelement.text               
                       vuln["<<vulnerability_evidences>>"] = []
                       subelement = subelement.find_next_sibling()
                   else:
                       vuln["<<"+key+">>"] = content   
                       data_structured['<<vulnerabilities>>'].append(vuln)
                      
                       
                       vuln = {} 
               elif subelement.name == 'h5':
                   subelement = subelement.find_next_sibling('h2') # Go to next vuln
                   if subelement is None:
                       break
               elif subelement.name == 'h6':                
                 
                   if not key:
                       key = subelement.text
                       subelement = subelement.find_next_sibling() 
                   else:
                       vuln["<<"+key+">>"] = content                   
                       key = subelement.text
                       content = ""
                       subelement = subelement.find_next_sibling()                        
               elif subelement.name == 'p':
                   if subelement.find_next_sibling().name == 'p':
                       content = content + subelement.text + '\r' 
                       print("1content value: "+ content)   
                       subelement = subelement.find_next_sibling() 
                   else:  
                       content = content + subelement.text
                       print("2content value: "+ content)   
                       subelement = subelement.find_next_sibling()
           vuln["<<"+key+">>"] = content            
           data_structured['<<vulnerabilities>>'].append(vuln) 
  
   count = 0
   count_vuln = 0
   for element in soup.find_all('h2'): 
       if element.find_previous_sibling('h1').text == "Vulnerabilities":  
           #print("Vulnerability: {}".format(element.text))
           next_element = element.find_next_sibling('h5')
           vulnerability_evidences = [] 

           while next_element.name == 'h5':
               evidence = {}
               vulnerability_evidence_note = "" 
               vulnerability_evidence_image_path = "" 
               content = ""
               #print("Evidence {}: {}".format(count,next_element.text))
               evidence_element = next_element.find_next_sibling(['h6','p'])
               
               while evidence_element.name == 'h6' or evidence_element.name == 'h2':
                   if  evidence_element.name == 'h2':
                       break
                   elif  evidence_element.name == 'h6' and evidence_element.text == "vulnerability_evidence_image_path":
                       sub_tag_p = evidence_element.find_next_sibling(['h6','p'])
                       content = ""
                       while sub_tag_p.name == 'p':
                           content = content + sub_tag_p.text + '\r'     
                           if sub_tag_p.find_next_sibling(['h6','p']) is None:
                               break
                           else:
                               sub_tag_p = sub_tag_p.find_next_sibling(['h6','p'])  
                       evidence["<<vulnerability_evidence_image_path>>"] = content
                   elif evidence_element.name == 'h6' and evidence_element.text == "vulnerability_evidence_note":
                       sub_tag_p = evidence_element.find_next_sibling(['h6','p'])
                       content = ""
                       while sub_tag_p.name == 'p':
                           content = content + sub_tag_p.text
                           if sub_tag_p.find_next_sibling(['h6','p']) is None:
                               break
                           else:
                               sub_tag_p = sub_tag_p.find_next_sibling(['h6','p']) 
                       evidence["<<vulnerability_evidence_note>>"] = content
                   if  evidence_element.name == 'h5' and evidence_element.text == "evidence":
                       #print("Agregando evidencia a la misma vulnerabilidad")
                       vulnerability_evidences.append(evidence)
                   
                   if evidence_element.find_next_sibling(['h6','h2','h5']) is not None:                
                       evidence_element = evidence_element.find_next_sibling(['h6','h2','h5']) 
                   else:
                       break
               
               vulnerability_evidences.append(evidence)
               #print("Evidence conent {}: {}".format(evidence["<<vulnerability_evidence_image_path>>"],evidence["<<vulnerability_evidence_note>>"]))
               
               if next_element.find_next_sibling(['h5','h2']) is not None: 
                   count = count + 1               
                   next_element = next_element.find_next_sibling(['h5','h2']) 
               else:
                   break
           data_structured['<<vulnerabilities>>'][count_vuln]["<<vulnerability_evidences>>"] = vulnerability_evidences
           count_vuln = count_vuln + 1
           
   json_data = json.dumps(data_structured, indent=2, separators=(',', ':'))
   
   return json_data

def removeLeadingZeros(str):
 
    # Regex to remove leading
    # zeros from a string
    regex = "^0+(?!$)"
 
    # Replaces the matched
    # value with given string
    str = re.sub(regex, "", str)
 
    print(str)


def decode_base64(data, altchars=b'+/'):
    """Decode base64, padding being optional.

    :param data: Base64 data as an ASCII byte string
    :returns: The decoded byte string.

    """
    data = re.sub(rb'[^a-zA-Z0-9%s]+' % altchars, b'', data)  # normalize
    missing_padding = len(data) % 4
    if missing_padding:
        data += b'='* (4 - missing_padding)
    return base64.b64decode(data, altchars)
