import win32com.client as win32
import win32clipboard as clip
import os
import json
import docx
import os.path
import numpy as np
from bs4 import BeautifulSoup
import pdb

from os.path import abspath
from win32com.client import constants
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor
from inspect import getmembers
import argparse
import logging

from datetime import datetime

import mistune
import xmltojson
import json
import requests
import html_to_json


#pdb.set_trace()



dt_string = datetime.now().strftime("%d-%m-%Y %H-%M-%S")
logging.basicConfig(filename='logs\\{} script.log'.format(dt_string), encoding='utf-8', level=logging.DEBUG)
logger = logging.getLogger(__name__)
visible_mode_win32com = False

def count_inrange(list1, l, r):
     
    # x for x in list1 is same as traversal in the list
    # the if condition checks for the number of numbers in the range
    # l to r
    # the return is stored in a list
    # whose length is the answer
    return len(list(x for x in list1 if l <= x <= r))
 

def merge_docx1(files, final_docx):
    # Start word application
    word = win32.gencache.EnsureDispatch("Word.Application")
    word.Visible = visible_mode_win32com
    word.DisplayAlerts = False

    # New blank document
    new_document = word.Documents.Add()
    for fn in files:
        # Open each file to be merged, copy the contents to the clipboard, and then close the file
        fn = abspath(fn)
        temp_document = word.Documents.Open(fn)
        word.Selection.WholeStory()
        word.Selection.Copy()
        temp_document.Close()
        # Paste to the end of the new document
        new_document.Range()
        word.Selection.Delete()
        word.Selection.Paste()
        
    clip.OpenClipboard()
    clip.EmptyClipboard()
    clip.CloseClipboard()
    # Save the final file and close the Word application
    new_document.SaveAs(final_docx)    
    new_document.Close(False)

    doc = docx.Document(final_docx)
    for table in doc.tables:
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="4C4C4C"/>'.format(nsdecls('w')))
        table.cell(0, 0)._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_2 = parse_xml(r'<w:shd {} w:fill="717171"/>'.format(nsdecls('w')))
        table.cell(0, 1)._tc.get_or_add_tcPr().append(shading_elm_2)
  
        if table.cell(0, 2).text != "-":
           risk_score_table = float(table.cell(0, 2).text)
        else:
           risk_score_table = 0    
        paragraph = table.cell(0, 2).paragraphs[0]
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        if(risk_score_table >= 0.1 and  risk_score_table <= 3.9):
            shading_elm_3 = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
            table.cell(0, 2)._tc.get_or_add_tcPr().append(shading_elm_3)
        elif(risk_score_table >= 4.0 and  risk_score_table <= 6.9):
            shading_elm_3 = parse_xml(r'<w:shd {} w:fill="FFC000"/>'.format(nsdecls('w')))
            table.cell(0, 2)._tc.get_or_add_tcPr().append(shading_elm_3)
        elif(risk_score_table >= 7.0 and  risk_score_table <= 8.9):
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255,255,255)
            shading_elm_3 = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w')))
            table.cell(0, 2)._tc.get_or_add_tcPr().append(shading_elm_3)
        elif(risk_score_table >= 9.0 and  risk_score_table <= 10): 
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255,255,255)           
            shading_elm_3 = parse_xml(r'<w:shd {} w:fill="C00000"/>'.format(nsdecls('w')))
            table.cell(0, 2)._tc.get_or_add_tcPr().append(shading_elm_3)
        elif(risk_score_table <= 0): 
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255,255,255)  
            shading_elm_3 = parse_xml(r'<w:shd {} w:fill="ADADAD"/>'.format(nsdecls('w')))
            table.cell(0, 2)._tc.get_or_add_tcPr().append(shading_elm_3)

           
    

            
        
        shading_elm_4 = parse_xml(r'<w:shd {} w:fill="717171"/>'.format(nsdecls('w')))
        table.cell(5, 0)._tc.get_or_add_tcPr().append(shading_elm_4)
        shading_elm_5 = parse_xml(r'<w:shd {} w:fill="717171"/>'.format(nsdecls('w')))
        table.cell(6, 0)._tc.get_or_add_tcPr().append(shading_elm_5)
        shading_elm_6 = parse_xml(r'<w:shd {} w:fill="717171"/>'.format(nsdecls('w')))
        table.cell(7, 0)._tc.get_or_add_tcPr().append(shading_elm_6)
        shading_elm_7 = parse_xml(r'<w:shd {} w:fill="717171"/>'.format(nsdecls('w')))
        table.cell(8, 0)._tc.get_or_add_tcPr().append(shading_elm_7)
        logging.info("Celda 0,0 tabla {}".format(table.cell(0, 0).text))
    doc.save(final_docx)

def merge_docx2(files, final_docx):
    output = wordapp.Documents.Add()
    output.Application.CutCopyMode = False
    for fn in files:
        output.Application.Selection.InsertFile(os.path.join(dn,fn) )        
    output.SaveAs(os.path.join(dn,'output.docx')  )
    output.Close(False)

dn = os.path.dirname(os.path.realpath(__file__))

parser = argparse.ArgumentParser(description='Web vulnerabilities report generator.')
parser.add_argument("-j", help='file JSON name', required=False)
parser.add_argument("-m", help='file JSON name', required=True)
args = vars(parser.parse_args())
analysis_filename = args['j']
markdown_filename = args['m']


# Create target directory & all intermediate directories if don't exists
if not os.path.isfile(markdown_filename):
    print('No existe ese archivo.')
    exit()

markdown_file  = open(markdown_filename,mode="r", encoding="utf-8") 
markdown_file_data = markdown_file.read()
markdown_file.close()
html_start = '<html><body>'
html_end = "</body></html>"
html_str = mistune.markdown(markdown_file_data)
complete_html = html_start+html_str+html_end

soup = BeautifulSoup(complete_html, "lxml")

invalid_tags = ['b','a', 'i', 'u']

for tag in invalid_tags: 
    for match in soup.findAll(tag):
        match.replaceWithChildren()

relevant_document_html_tags = ['h1','h2','h3','h4','h5','h6','p']
vuln_tags = ['h2','h6','p']
data_structured = {}


for element in soup.find_all('h1'):  
    if  element.text == 'App information':  
        tag = element.find_next_sibling(relevant_document_html_tags)
        while tag.name == 'h6' or tag.name == 'p':
            content = "" 
            if tag.name == 'h6': 
                key = tag.text                
            elif tag.name == 'p':
                subtag = tag                
                while subtag.name == 'p':
                    content = content + subtag.text 
                    subtag = subtag.find_next_sibling(relevant_document_html_tags)
            data_structured["<<"+key+">>"] = content
            tag = tag.find_next_sibling(relevant_document_html_tags)


data_structured['<<vulnerabilities>>'] = []

for element in soup.find_all('h1'):  
    if  element.text == 'Vulnerabilities':        
        vuln = {}    
        key = ""
        tag = element.find_next_sibling(vuln_tags)
        while tag.name == 'h6' or tag.name == 'p' or tag.name == 'h2':
            if tag.name == 'h2':
                
                vuln = {}  
                vuln["<<name>>"] = tag.text               
                vuln["<<evidences>>"] = []
                tag_arg = tag.find_next_sibling(['h6','p','h5']) 
                while tag_arg.name == 'h6' or tag_arg.name == 'p' or tag_arg.name == 'h5':
                    #print(tag_arg.text)
                    #print(tag_arg.name)
                    if tag_arg.name == 'h5':
                        break
                    elif tag_arg.name == 'h6': 
                        key = tag_arg.text
                        content = ""
                    elif tag_arg.name == 'p':  
                        subtag = tag_arg                             
                        while subtag.name == 'p':
                            content = content + subtag.text                     
                            if subtag.find_next_sibling(['h6','p','h5']) is None:
                                break
                            else:
                                subtag = subtag.find_next_sibling(['h6','p','h5']) 
                    vuln["<<"+key+">>"] = content
                    if tag_arg.find_next_sibling(['h6','p','h5']) is None:
                        break
                    else:                        
                        tag_arg = tag_arg.find_next_sibling(['h6','p','h5'])               
                data_structured['<<vulnerabilities>>'].append(vuln)
            if tag.find_next_sibling(['h2']) is not None:
                    tag = tag.find_next_sibling(['h2'])
            else:
                break


count = 0
count_vuln = 0
for element in soup.find_all('h2'): 
    if element.find_previous_sibling('h1').text == "Vulnerabilities":  
        #print("Vulnerability: {}".format(element.text))
        next_element = element.find_next_sibling('h5')
        evidences = [] 
        
        while next_element.name == 'h5':
            evidence = {}
            evidence_note = "" 
            evidence_image_path = "" 
            content = ""
            #print("Evidence {}: {}".format(count,next_element.text))
            evidence_element = next_element.find_next_sibling(['h6','p'])
            
            while evidence_element.name == 'h6' or evidence_element.name == 'h2':
                if  evidence_element.name == 'h2':
                    break
                if  evidence_element.name == 'h6' and evidence_element.text == "evidence_image_path":
                    sub_tag_p = evidence_element.find_next_sibling(['h6','p'])
                    content = ""
                    while sub_tag_p.name == 'p':
                        content = content + sub_tag_p.text
                        if sub_tag_p.find_next_sibling(['h6','p']) is None:
                            break
                        else:
                            sub_tag_p = sub_tag_p.find_next_sibling(['h6','p'])  
                    evidence["<<evidence_image_path>>"] = content
                elif evidence_element.name == 'h6' and evidence_element.text == "evidence_note":
                    sub_tag_p = evidence_element.find_next_sibling(['h6','p'])
                    content = ""
                    while sub_tag_p.name == 'p':
                        content = content + sub_tag_p.text
                        if sub_tag_p.find_next_sibling(['h6','p']) is None:
                            break
                        else:
                            sub_tag_p = sub_tag_p.find_next_sibling(['h6','p']) 
                    evidence["<<evidence_note>>"] = content
                
                if evidence_element.find_next_sibling(['h6','h2']) is not None:                
                    evidence_element = evidence_element.find_next_sibling(['h6','h2']) 
                else:
                    break
            evidences.append(evidence)
            if next_element.find_next_sibling(['h5','h2']) is not None: 
                count = count + 1               
                next_element = next_element.find_next_sibling(['h5','h2']) 
            else:
                break
        data_structured['<<vulnerabilities>>'][count_vuln]["<<evidences>>"] = evidences
        count_vuln = count_vuln + 1
        
json_data = json.dumps(data_structured, indent=2, separators=(',', ':'))

analysis_filename = 'temp.json'
with open( os.path.join(dn,analysis_filename), 'w') as f:
    json.dump(json_data, f)

# Opening JSON file
with open(analysis_filename, encoding='utf-8') as json_file:
    data = json.loads(json.load(json_file))

path_file_name = os.path.join(dn,data['<<template_name>>'])

name_file = data['<<analysis_id>>'] + ' ' + data ['<<name_app>>'] + ' - ' + data['<<analysis_version_format_01>>'] +".docx"
name_file = name_file.replace("/", "-")

print(name_file)
print(name_file)
print(name_file)
print(name_file)

wordapp = win32.gencache.EnsureDispatch("Word.Application")
wordapp.Visible = visible_mode_win32com
wordapp.DisplayAlerts = False

doc = wordapp.Documents.Open(path_file_name)
doc.Activate()

wordapp.ActiveDocument.TrackRevisions = False  # Maybe not need this (not really but why not)
wordapp.Selection.GoTo(win32.constants.wdGoToPage, win32.constants.wdGoToAbsolute, "2")



for From in data.keys():
    logging.info(From)
    if (From != '<<executive_resume>>'):
        wordapp.ActiveWindow.ActivePane.View.SeekView =win32.constants.wdSeekMainDocument
        wordapp.Selection.Find.Execute(From, False, False, False, False, False, True, win32.constants.wdFindContinue, False, data[From], win32.constants.wdReplaceAll) 
        wordapp.ActiveWindow.ActivePane.View.SeekView = win32.constants.wdSeekCurrentPageHeader
        wordapp.Selection.Find.Execute(From, False, False, False, False, False, True, win32.constants.wdFindContinue, False, data[From], win32.constants.wdReplaceAll) 

wordapp.Selection.GoTo(win32.constants.wdGoToPage, win32.constants.wdGoToAbsolute, "2")
wordapp.ActiveDocument.SaveAs(os.path.join(dn,name_file))
doc.Close(False)

count = 1

vulnerabilities_tables = []

sorted_asc_vulns = data['<<vulnerabilities>>']
sorted_asc_vulns.sort(key=lambda x: float(x["<<risk_score>>"]),reverse = True)
for vunl in sorted_asc_vulns:
    # Upper Case letters
    vunl['<<name_upper>>'] = vunl['<<name>>'].upper()

for vunl in sorted_asc_vulns:
    logging.info(vunl)
    table_template = os.path.join(dn,'templates\\Plantilla de tabla de vulnerabilidades.docx')
    wordapp = win32.gencache.EnsureDispatch("Word.Application")
    wordapp.Visible = visible_mode_win32com
    wordapp.DisplayAlerts = False
    doc = wordapp.Documents.Open(table_template )
    doc.Activate()
    wordapp.Selection.HomeKey(Unit=win32.constants.wdStory)   
    for From in vunl.keys():
        word_replace = vunl[From]
        if word_replace == "0":
           word_replace = "-"
        wordapp.Selection.HomeKey(Unit=win32.constants.wdStory)
        try:
            wordapp.Selection.Find.Execute(From) 
            wordapp.Selection.Text = word_replace
        except Exception as e: 
            print(e)
    wordapp.Selection.EndKey(Unit=win32.constants.wdStory)
    
    try:
        for vunl in vunl['<<evidences>>']:
            shape = wordapp.Selection.InlineShapes.AddPicture(FileName=vunl['<<evidence_image_path>>'],LinkToFile=False, SaveWithDocument=True )
            shape.LockAspectRatio = True
            #shape.Width = 450; # Change width works 
            wordapp.Selection.TypeText(Text='\n\r'+vunl['<<evidence_note>>']+'\n\r')
        if (count  <  len(data['<<vulnerabilities>>'])):
            wordapp.Selection.InsertBreak(Type=win32.constants.wdPageBreak)  
    except Exception as e: 
        print(e)

    vulnerabilities_tables.append('Table_'+ str(count)+ ".docx")
    wordapp.ActiveDocument.SaveAs(os.path.join(dn,'Table_'+ str(count)+ ".docx"))
    doc.Close(False)
    count = count + 1


vulns_table_file_name = 'Vulnerabilities -'+ data['<<analysis_id>>'] + ' ' + data ['<<name_app>>'] + ' - ' + data['<<analysis_version_format_01>>'] + ".docx"
vulns_table_file_name = vulns_table_file_name.replace("/", "-")

merge_docx1(vulnerabilities_tables,os.path.join(dn,vulns_table_file_name))

doca = wordapp.Documents.Open(os.path.join(dn,name_file))
doca.Activate()

wordapp.Selection.Find.Execute('<<vulnerabilities_tables>>') 
wordapp.Selection.InsertFile(FileName=os.path.join(dn,vulns_table_file_name), Range="", ConfirmConversions=False, Link=False, Attachment=False)
wordapp.Selection.InsertBreak(Type=win32.constants.wdPageBreak)  
# Go to start document
wordapp.Selection.HomeKey(Unit=win32.constants.wdStory)
wordapp.Selection.Find.Execute('<<executive_resume>>') 
wordapp.Selection.Text = data['<<executive_resume>>']
wordapp.Selection.HomeKey(Unit=win32.constants.wdStory)
wordapp.Selection.Find.Execute('<<level_max>>') 
wordapp.Selection.Font.Bold = True


risk_scores = []
risk_list = []
remediation_list = []
for vunl in data['<<vulnerabilities>>']:      
    risk_scores.append(float(vunl["<<risk_score>>"]))
    risk_list.append(vunl["<<risk>>"]) 
    remediation_list.append(vunl["<<remediation>>"]) 

# Dictionary count
dict_of_counts = {}
dict_of_counts["INFORMATIVA"] = count_inrange(risk_scores,-1,0)
dict_of_counts["BAJA"] = count_inrange(risk_scores,0.1,3.9)
dict_of_counts["MEDIA"] = count_inrange(risk_scores,4.0,6.9)
dict_of_counts["ALTA"] = count_inrange(risk_scores,7.0,8.9)
dict_of_counts["CRÍTICA"] = count_inrange(risk_scores,9.0,10)


max_level = "BAJO"
# Have to identify the index of the graph you want to handle
if (wordapp.ActiveDocument.InlineShapes(1).Type == 12): # Is a chart
    wb = wordapp.ActiveDocument.InlineShapes(1).Chart.ChartData.Workbook
    wb.Application.Visible = visible_mode_win32com
    #wordapp.ActiveDocument.InlineShapes(1).Chart.ChartData.Activate()
    SourceSheet = wb.ActiveSheet

    if (dict_of_counts["INFORMATIVA"] > 0):
        SourceSheet.Range("B2").Value2 = dict_of_counts["INFORMATIVA"]
        max_level = "INFORMATIVO"
    else:
        wordapp.ActiveDocument.InlineShapes(1).Chart.ChartGroups(1).FullCategoryCollection(1).IsFiltered = True
    if (dict_of_counts["BAJA"] > 0):
        SourceSheet.Range("B3").Value2 = dict_of_counts["BAJA"]
        max_level = "BAJO"
    else:
        wordapp.ActiveDocument.InlineShapes(1).Chart.ChartGroups(1).FullCategoryCollection(2).IsFiltered = True
    if (dict_of_counts["MEDIA"] > 0):
        SourceSheet.Range("B4").Value2 = dict_of_counts["MEDIA"]
        max_level = "MEDIO"
    else:
        wordapp.ActiveDocument.InlineShapes(1).Chart.ChartGroups(1).FullCategoryCollection(3).IsFiltered = True
    if (dict_of_counts["ALTA"] > 0):
        SourceSheet.Range("B5").Value2 = dict_of_counts["ALTA"]
        max_level = "ALTO"
    else:
        wordapp.ActiveDocument.InlineShapes(1).Chart.ChartGroups(1).FullCategoryCollection(4).IsFiltered = True
    if (dict_of_counts["CRÍTICA"] > 0):
        SourceSheet.Range("B6").Value2 = dict_of_counts["CRÍTICA"]
        max_level = "CRÍTICO"
    else:
        wordapp.ActiveDocument.InlineShapes(1).Chart.ChartGroups(1).FullCategoryCollection(5).IsFiltered = True
    
    # xlCellTypeBlanks =  4
    #SourceSheet.Range("A2:B5").SpecialCells(4).Delete()

    wb.Close()
    wordapp.ActiveDocument.InlineShapes(1).Chart.Refresh
    #wordapp.ActiveDocument.InlineShapes(1).Chart.ChartGroups(1).FullCategoryCollection(1).IsFiltered = True
    #wordapp.ActiveDocument.InlineShapes(1).Chart.SeriesCollection(1).DataLabels.ShowValue = False
    
wordapp.Selection.Find.Execute("<<level_max>>", False, False, False, False, False, True, win32.constants.wdFindContinue, False, max_level, win32.constants.wdReplaceAll) 

# Get the correct index of table
#print(doca.Tables)

logging.info(doca.Tables(2).Cell(1, 1).Range.Text)
logging.info(doca.Tables(2).Cell(3, 1).Range.Text)
logging.info(doca.Tables(2).Rows.Count)

doca.Tables(2).Rows.Add()
doca.Tables(2).Cell(3, 1).Select() 
wordapp.Selection.SelectRow() 
wordapp.Selection.Cells.Delete(ShiftCells=win32.constants.wdDeleteCellsEntireRow)

#doca.Tables(2).Cell(3, 2).Delete()
#doca.Tables(2).Cell(3, 1).Delete()

for vunl in sorted_asc_vulns:
    if(float(vunl['<<risk_score>>']) >= 0.1 and  float(vunl['<<risk_score>>']) <= 3.9):
        vunl['<<level_risk>>'] = 'BAJA'
        vunl['<<level_risk_olecolor>>'] = 65535
        vunl['<<level_risk_text_olecolor>>'] = -16777216
        vunl['<<level_risk_text_rgbcolor>>'] = RGBColor(255,255,255)
    elif(float(vunl['<<risk_score>>']) >= 4.0 and  float(vunl['<<risk_score>>']) <= 6.9):
        vunl['<<level_risk>>'] = 'MEDIA'
        vunl['<<level_risk_olecolor>>'] = 49407
        vunl['<<level_risk_text_olecolor>>'] = -16777216
        vunl['<<level_risk_text_rgbcolor>>'] = RGBColor(255,255,255)
    elif(float(vunl['<<risk_score>>']) >= 7.0 and  float(vunl['<<risk_score>>']) <= 8.9):
        vunl['<<level_risk>>'] = 'ALTA'
        vunl['<<level_risk_olecolor>>'] = 255
        vunl['<<level_risk_text_olecolor>>'] = 16777215
        vunl['<<level_risk_text_rgbcolor>>'] = RGBColor(0,0,0)
    elif(float(vunl['<<risk_score>>']) >= 9.0 and  float(vunl['<<risk_score>>']) <= 10): 
        vunl['<<level_risk>>'] = 'CRÍTICA'
        vunl['<<level_risk_olecolor>>'] = 192
        vunl['<<level_risk_text_olecolor>>'] = 16777215
        vunl['<<level_risk_text_rgbcolor>>'] =  RGBColor(0,0,0)
    elif(float(vunl['<<risk_score>>']) <= 0): 
        vunl['<<level_risk>>'] = 'INFORMATIVA'
        vunl['<<level_risk_olecolor>>'] = 11382189
        vunl['<<level_risk_text_olecolor>>'] = -16777216
        vunl['<<level_risk_text_rgbcolor>>'] = RGBColor(255,255,255)
        
    
index = 3
for vunl in sorted_asc_vulns:
    doca.Tables(2).Cell(index, 1).Range.Text = vunl['<<level_risk>>']
    doca.Tables(2).Cell(index, 1).Shading.BackgroundPatternColor = vunl['<<level_risk_olecolor>>']
    doca.Tables(2).Cell(index, 1).Range.Font.Color = vunl['<<level_risk_text_olecolor>>']
    doca.Tables(2).Cell(index, 2).Range.Text =  "{0:3}. {1}".format(index-2,vunl['<<name>>']) 
    
    index = index + 1 
    doca.Tables(2).Rows.Add()
doca.Tables(2).Cell(index, 1).Select() 
wordapp.Selection.SelectRow() 
wordapp.Selection.Cells.Delete(ShiftCells=win32.constants.wdDeleteCellsEntireRow)


wordapp.Selection.HomeKey(Unit=win32.constants.wdStory)
wordapp.Selection.Find.Execute('<<risk_list>>') 
wordapp.Selection.Text = '\n'.join(risk_list)

wordapp.Selection.HomeKey(Unit=win32.constants.wdStory)
wordapp.Selection.Find.Execute('<<recomendation_list>>') 
wordapp.Selection.Text = '\n'.join(remediation_list)

# Delete rows 

#doca.Tables(2).Rows(3).Delete()
#doca.Tables(2).Rows(3).EntireRow.Delete()
#doca.Tables(2).Rows(3).Delete()
#for vunl in sorted_asc_vulns:
#    doca.Tables(2).Rows.Add()
#    sorted_asc_vulns['']


#doca.Tables(2).Cell(4, 1).Range.Text = "BAJA"
#doca.Tables(2).Cell(4, 2).Range.Text = "2. HOLA HOLA"

#doca.Tables(1).Range.Cell(0,0).Interior.Color = rgbToInt((255,0,0))
#print(doca.Tables(1).Cells.Count)
#print(doca.Tables(1).Range.Rows)
#print(doca.Tables(1).columns)
# Columna, Fila

#print(doca.Tables(2).Cell(1, 1).Range.Text)
#print(doca.Tables(2).Cell(1, 2).Range.Text)

#print(doca.Tables(1).Cell(0, 0).Range.Text)
#doca.Tables(1).Rows(1).Delete()

doca.TablesOfContents(1).Update()
wordapp.ActiveDocument.Save()
doca.Close(SaveChanges=True)

wordapp.Application.Quit()

for f in vulnerabilities_tables:
    # Delete temporal files generated
    if os.path.isfile(f): # this makes the code more robust
            os.remove(f)

if os.path.isfile(vulns_table_file_name): # this makes the code more robust
            os.remove(vulns_table_file_name)
