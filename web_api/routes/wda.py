import xmltojson
import win32com.client as win32
import win32com
import win32clipboard as clip
import sys
import sys
import sys
import shutil
import requests
import re
import pythoncom
import pyfiglet
import pdb
import os.path
import os.path
import os
import os
import os
import numpy as np
import mistune
import logging
import locale
import json
import json
import json
import html_to_json
import docx
import base64
import argparse
from win32com.client import constants
from schemas.wda import serializeDict, serializeList, wdaDict
from os.path import abspath
from models.wda import VulnerabilityAnalisis, Vulnerability, Evidence
from inspect import getmembers
from fastapi.responses import HTMLResponse
from fastapi.responses import FileResponse
from fastapi import APIRouter
from fastapi import Request
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx import Document
from datetime import datetime, timedelta
from config.db import db 
from bson import ObjectId
from bs4 import BeautifulSoup



#dn = os.path.dirname(os.path.abspath(sys.argv[0]))
dn = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop','avums') 

dt_string = datetime.now().strftime("%d-%m-%Y %H-%M-%S")
visible_mode_win32com = True



wda = APIRouter() 



class bcolors:
    HEADER = '\033[95m'
    OKBLUE = '\033[94m'
    OKCYAN = '\033[96m'
    OKGREEN = '\033[92m'
    WARNING = '\033[93m'
    FAIL = '\033[91m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'

#pdb.set_trace()



def covx_to_pdf(infile, outfile, word):
    """Convert a Word .docx to PDF"""
    wdFormatPDF = 17
    doc = word.Documents.Open(infile)
    doc.SaveAs(outfile, FileFormat=wdFormatPDF)
    doc.Close()

def count_inrange(list1, l, r):
     
    # x for x in list1 is same as traversal in the list
    # the if condition checks for the number of numbers in the range
    # l to r
    # the return is stored in a list
    # whose length is the answer
    return len(list(x for x in list1 if l <= x <= r))
 

def merge_docx1(files, final_docx_name):
    # Start word application
    word = win32.gencache.EnsureDispatch("Word.Application")
    word.Visible = visible_mode_win32com
    word.DisplayAlerts = False

    # New blank document
    new_document = word.Documents.Add()
    for fn in files:
        # Open each file to be merged, copy the contents to the clipboard, and then close the file
        fn = abspath(fn)
        temp_document = word.Documents.Open(fn)
        word.Selection.WholeStory()
        word.Selection.Copy()
        temp_document.Close()
        # Paste to the end of the new document
        new_document.Range()
        word.Selection.Delete()
        word.Selection.Paste()        
    clip.OpenClipboard()
    clip.EmptyClipboard()
    clip.CloseClipboard()
    # Save the final file and close the Word application
    new_document.SaveAs(os.path.join(dn,"outputs",final_docx_name))    
    new_document.Close(False)

    doc = docx.Document(os.path.join(dn,"outputs",final_docx_name))
    for table in doc.tables:
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="4C4C4C"/>'.format(nsdecls('w')))
        table.cell(0, 0)._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_2 = parse_xml(r'<w:shd {} w:fill="717171"/>'.format(nsdecls('w')))
        table.cell(0, 1)._tc.get_or_add_tcPr().append(shading_elm_2)
  
        if table.cell(0, 2).text != "-":
           risk_score_table = float(table.cell(0, 2).text)
        else:
           risk_score_table = 0    
        paragraph = table.cell(0, 2).paragraphs[0]
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        if(risk_score_table >= 0.1 and  risk_score_table <= 3.9):
            shading_elm_3 = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
            table.cell(0, 2)._tc.get_or_add_tcPr().append(shading_elm_3)
        elif(risk_score_table >= 4.0 and  risk_score_table <= 6.9):
            shading_elm_3 = parse_xml(r'<w:shd {} w:fill="FFC000"/>'.format(nsdecls('w')))
            table.cell(0, 2)._tc.get_or_add_tcPr().append(shading_elm_3)
        elif(risk_score_table >= 7.0 and  risk_score_table <= 8.9):
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255,255,255)
            shading_elm_3 = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w')))
            table.cell(0, 2)._tc.get_or_add_tcPr().append(shading_elm_3)
        elif(risk_score_table >= 9.0 and  risk_score_table <= 10): 
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255,255,255)           
            shading_elm_3 = parse_xml(r'<w:shd {} w:fill="C00000"/>'.format(nsdecls('w')))
            table.cell(0, 2)._tc.get_or_add_tcPr().append(shading_elm_3)
        elif(risk_score_table <= 0): 
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255,255,255)  
            shading_elm_3 = parse_xml(r'<w:shd {} w:fill="ADADAD"/>'.format(nsdecls('w')))
            table.cell(0, 2)._tc.get_or_add_tcPr().append(shading_elm_3)
            
        
        shading_elm_4 = parse_xml(r'<w:shd {} w:fill="717171"/>'.format(nsdecls('w')))
        table.cell(5, 0)._tc.get_or_add_tcPr().append(shading_elm_4)
        shading_elm_5 = parse_xml(r'<w:shd {} w:fill="717171"/>'.format(nsdecls('w')))
        table.cell(6, 0)._tc.get_or_add_tcPr().append(shading_elm_5)
        shading_elm_6 = parse_xml(r'<w:shd {} w:fill="717171"/>'.format(nsdecls('w')))
        table.cell(7, 0)._tc.get_or_add_tcPr().append(shading_elm_6)
        shading_elm_7 = parse_xml(r'<w:shd {} w:fill="717171"/>'.format(nsdecls('w')))
        table.cell(8, 0)._tc.get_or_add_tcPr().append(shading_elm_7)
        logging.info("Celda 0,0 tabla {}".format(table.cell(0, 0).text))
    doc.save(os.path.join(dn,"outputs",final_docx_name))



def generate_report(data):
   try:
      shutil.rmtree(win32com.__gen_path__, ignore_errors=True) 
   except Exception as e: 
       print(e)
   # Opening JSON file
   #with open(analysis_filename, encoding='utf-8') as json_file:
   #    data = json.loads(json.load(json_file))
   
   #template_file_path = os.path.join(dn,'templates',data['<<template_name>>'])
   
   template_file_path = os.path.join(dn,'templates',data['<<template_name>>']).replace('\r', '')
   
   name_file = data['<<analysis_id>>'] + ' ' + data ['<<name_app>>'] + ' - ' + data['<<analysis_version_format_01>>'] + ".docx"
   name_file = name_file.replace("/", "-").replace('\r', '')
   
   wordapp = win32.gencache.EnsureDispatch("Word.Application",pythoncom.CoInitialize())
   wordapp.Visible = visible_mode_win32com
   wordapp.DisplayAlerts = False
   
   doc = wordapp.Documents.Open(template_file_path)
   doc.Activate()
   
   wordapp.ActiveDocument.TrackRevisions = False  # Maybe not need this (not really but why not)
   wordapp.Selection.GoTo(win32.constants.wdGoToPage, win32.constants.wdGoToAbsolute, "2")
   
   try:
      for From in data.keys():
          logging.info(From)
          if (From != '<<executive_resume>>'):
              wordapp.ActiveWindow.ActivePane.View.SeekView =win32.constants.wdSeekMainDocument
              wordapp.Selection.Find.Execute(From, False, False, False, False, False, True, win32.constants.wdFindContinue, False, data[From], win32.constants.wdReplaceAll) 
              wordapp.ActiveWindow.ActivePane.View.SeekView = win32.constants.wdSeekCurrentPageHeader              
              wordapp.Selection.Find.Execute(From, False, False, False, False, False, True, win32.constants.wdFindContinue, False, data[From], win32.constants.wdReplaceAll) 
          wordapp.ActiveWindow.ActivePane.View.SeekView =win32.constants.wdSeekMainDocument
   except Exception as e: 
           print(e)  

 
   wordapp.Selection.GoTo(win32.constants.wdGoToPage, win32.constants.wdGoToAbsolute, "2")
   wordapp.ActiveDocument.SaveAs(os.path.join(dn,"outputs",name_file))

   doc.TablesOfContents(1).Update()
   wordapp.ActiveDocument.Save()
   doc.Close(SaveChanges=True)


   if data['<<vulnerabilities>>']:
        try:
            count = 1
            
            vulnerabilities_tables = []
            
            sorted_asc_vulns = data['<<vulnerabilities>>']

            sorted_asc_vulns.sort(key=lambda x: float(x["<<vulnerability_risk_score>>"]),reverse = True)
 

            for vunl in sorted_asc_vulns:
                # Upper Case letters
                vunl['<<vulnerability_name_upper>>'] = vunl['<<vulnerability_name>>'].upper()
            
            for vunl in sorted_asc_vulns:
              
                table_template = os.path.join(dn,'templates','template-vulns-table.docx')
                
                wordapp = win32.gencache.EnsureDispatch("Word.Application")
                wordapp.Visible = visible_mode_win32com
                wordapp.DisplayAlerts = False
                doc = wordapp.Documents.Open(table_template)
                doc.Activate()
                wordapp.Selection.HomeKey(Unit=win32.constants.wdStory)   
                for From in vunl.keys():                   
                    word_replace = vunl[From]
                    if word_replace == "0":
                       word_replace = "-"
                    wordapp.Selection.HomeKey(Unit=win32.constants.wdStory)
                    if From != "<<vulnerability_evidence_note>>" and From !="<<vulnerability_evidence_image_path>>" and From !="<<vulnerability_evidences>>":
                        try:
                            wordapp.Selection.Find.Execute(From) 
                            wordapp.Selection.Text = word_replace
                        except Exception as e: 
                            print(e)
         
               
                
         
                try:   
                    for evidence in vunl['<<vulnerability_evidences>>'] :  
                          wordapp.Selection.EndKey(Unit=win32.constants.wdStory)                                          
                          image_path = evidence['<<vulnerability_evidence_image_path>>']                          
                          shape = wordapp.Selection.InlineShapes.AddPicture(FileName=image_path,LinkToFile=False, SaveWithDocument=True )
                          shape.LockAspectRatio = True
                          #shape.Width = 450; # Change width works 
                          wordapp.Selection.TypeText(Text='\n\r'+evidence['<<vulnerability_evidence_note>>']+'\n\r')
                          wordapp.ActiveDocument.Paragraphs.Last.Alignment = win32.constants.wdAlignParagraphJustify
                          wordapp.Selection.TypeParagraph
                                                     
                except Exception as e: 
                    print(e)


                if (count  <  len(data['<<vulnerabilities>>'])):
                    wordapp.Selection.EndKey(Unit=win32.constants.wdStory) 
                    wordapp.Selection.InsertBreak(Type=win32.constants.wdPageBreak)  
            
                vulnerabilities_tables.append(os.path.join(dn,"outputs",'Table_'+ str(count)+ ".docx"))
                wordapp.ActiveDocument.SaveAs(os.path.join(dn,"outputs",'Table_'+ str(count)+ ".docx"))
                doc.Close(False)
                count = count + 1
            
            vulns_table_file_name =   'Vulnerabilities {} {} {}{}'.format(data['<<analysis_id>>'],data ['<<name_app>>'],data['<<analysis_version_format_01>>'],".docx")
            vulns_table_file_name = vulns_table_file_name.replace("/", "-").replace('\r', '')
            full_vulns_table_file_name = os.path.join(dn,"outputs",vulns_table_file_name).replace('\r', '')
            
            #print(vulns_table_file_name)
            
            merge_docx1(vulnerabilities_tables,vulns_table_file_name)

            
            
            doca = wordapp.Documents.Open(os.path.join(dn,"outputs",name_file))
            doca.Activate()
            
            wordapp.Selection.Find.Execute('<<vulnerabilities_tables>>') 
            wordapp.Selection.InsertFile(FileName=full_vulns_table_file_name, Range="", ConfirmConversions=False, Link=False, Attachment=False)
            wordapp.Selection.InsertBreak(Type=win32.constants.wdPageBreak)  
            # Go to start document
            wordapp.Selection.HomeKey(Unit=win32.constants.wdStory)
            wordapp.Selection.Find.Execute('<<executive_resume>>') 
            wordapp.Selection.Text = data['<<executive_resume>>']
            wordapp.Selection.HomeKey(Unit=win32.constants.wdStory)
            wordapp.Selection.Find.Execute('<<level_max>>') 
            wordapp.Selection.Font.Bold = True            
            
            risk_scores = []
            risk_list = []
            remediation_list = []
            for vunl in data['<<vulnerabilities>>']:      
                risk_scores.append(float(vunl["<<vulnerability_risk_score>>"]))
                risk_list.append(vunl["<<vulnerability_risk>>"]) 
                remediation_list.append(vunl["<<vulnerability_remediation>>"]) 

            
            
            # Dictionary count
            dict_of_counts = {}
            dict_of_counts["INFORMATIVA"] = count_inrange(risk_scores,-1,0)
            dict_of_counts["BAJA"] = count_inrange(risk_scores,0.1,3.9)
            dict_of_counts["MEDIA"] = count_inrange(risk_scores,4.0,6.9)
            dict_of_counts["ALTA"] = count_inrange(risk_scores,7.0,8.9)
            dict_of_counts["CRÍTICA"] = count_inrange(risk_scores,9.0,10)
            
            
            max_level = "BAJO"
    
            # Have to identify the index of the graph you want to handle
            if (wordapp.ActiveDocument.InlineShapes(1).Type == 12): # Is a chart
                chart_wb = wordapp.ActiveDocument.InlineShapes(1).Chart.ChartData.Workbook
                chart_wb.Application.Visible = visible_mode_win32com
                #wordapp.ActiveDocument.InlineShapes(1).Chart.ChartData.Activate()
                SourceSheet = chart_wb.ActiveSheet
            
                if (dict_of_counts["INFORMATIVA"] > 0):
                    SourceSheet.Range("B2").Value2 = dict_of_counts["INFORMATIVA"]
                    max_level = "INFORMATIVO"
                else:
                    wordapp.ActiveDocument.InlineShapes(1).Chart.ChartGroups(1).FullCategoryCollection(1).IsFiltered = True
                if (dict_of_counts["BAJA"] > 0):
                    SourceSheet.Range("B3").Value2 = dict_of_counts["BAJA"]
                    max_level = "BAJO"
                else:
                    wordapp.ActiveDocument.InlineShapes(1).Chart.ChartGroups(1).FullCategoryCollection(2).IsFiltered = True
                if (dict_of_counts["MEDIA"] > 0):
                    SourceSheet.Range("B4").Value2 = dict_of_counts["MEDIA"]
                    max_level = "MEDIO"
                else:
                    wordapp.ActiveDocument.InlineShapes(1).Chart.ChartGroups(1).FullCategoryCollection(3).IsFiltered = True
                if (dict_of_counts["ALTA"] > 0):
                    SourceSheet.Range("B5").Value2 = dict_of_counts["ALTA"]
                    max_level = "ALTO"
                else:
                    wordapp.ActiveDocument.InlineShapes(1).Chart.ChartGroups(1).FullCategoryCollection(4).IsFiltered = True
                if (dict_of_counts["CRÍTICA"] > 0):
                    SourceSheet.Range("B6").Value2 = dict_of_counts["CRÍTICA"]
                    max_level = "CRÍTICO"
                else:
                    wordapp.ActiveDocument.InlineShapes(1).Chart.ChartGroups(1).FullCategoryCollection(5).IsFiltered = True
                
                # xlCellTypeBlanks =  4
                #SourceSheet.Range("A2:B5").SpecialCells(4).Delete()
            
                chart_wb.Close()
                wordapp.ActiveDocument.InlineShapes(1).Chart.Refresh
                #wordapp.ActiveDocument.InlineShapes(1).Chart.ChartGroups(1).FullCategoryCollection(1).IsFiltered = True
                #wordapp.ActiveDocument.InlineShapes(1).Chart.SeriesCollection(1).DataLabels.ShowValue = False
                
            wordapp.Selection.Find.Execute("<<level_max>>", False, False, False, False, False, True, win32.constants.wdFindContinue, False, max_level, win32.constants.wdReplaceAll) 
            
            # Get the correct index of table
            #print(doca.Tables)

            
            logging.info(doca.Tables(2).Cell(1, 1).Range.Text)
            logging.info(doca.Tables(2).Cell(3, 1).Range.Text)
            logging.info(doca.Tables(2).Rows.Count)
            
            doca.Tables(2).Rows.Add()
            doca.Tables(2).Cell(3, 1).Select() 
            wordapp.Selection.SelectRow() 
            wordapp.Selection.Cells.Delete(ShiftCells=win32.constants.wdDeleteCellsEntireRow)
            
            #doca.Tables(2).Cell(3, 2).Delete()
            #doca.Tables(2).Cell(3, 1).Delete()
            
            for vunl in sorted_asc_vulns:
                if(float(vunl['<<vulnerability_risk_score>>']) >= 0.1 and  float(vunl['<<vulnerability_risk_score>>']) <= 3.9):
                    vunl['<<level_risk>>'] = 'BAJA'
                    vunl['<<level_risk_olecolor>>'] = 65535
                    vunl['<<level_risk_text_olecolor>>'] = -16777216
                    vunl['<<level_risk_text_rgbcolor>>'] = RGBColor(255,255,255)
                elif(float(vunl['<<vulnerability_risk_score>>']) >= 4.0 and  float(vunl['<<vulnerability_risk_score>>']) <= 6.9):
                    vunl['<<level_risk>>'] = 'MEDIA'
                    vunl['<<level_risk_olecolor>>'] = 49407
                    vunl['<<level_risk_text_olecolor>>'] = -16777216
                    vunl['<<level_risk_text_rgbcolor>>'] = RGBColor(255,255,255)
                elif(float(vunl['<<vulnerability_risk_score>>']) >= 7.0 and  float(vunl['<<vulnerability_risk_score>>']) <= 8.9):
                    vunl['<<level_risk>>'] = 'ALTA'
                    vunl['<<level_risk_olecolor>>'] = 255
                    vunl['<<level_risk_text_olecolor>>'] = 16777215
                    vunl['<<level_risk_text_rgbcolor>>'] = RGBColor(0,0,0)
                elif(float(vunl['<<vulnerability_risk_score>>']) >= 9.0 and  float(vunl['<<vulnerability_risk_score>>']) <= 10): 
                    vunl['<<level_risk>>'] = 'CRÍTICA'
                    vunl['<<level_risk_olecolor>>'] = 192
                    vunl['<<level_risk_text_olecolor>>'] = 16777215
                    vunl['<<level_risk_text_rgbcolor>>'] =  RGBColor(0,0,0)
                elif(float(vunl['<<vulnerability_risk_score>>']) <= 0): 
                    vunl['<<level_risk>>'] = 'INFORMATIVA'
                    vunl['<<level_risk_olecolor>>'] = 11382189
                    vunl['<<level_risk_text_olecolor>>'] = -16777216
                    vunl['<<level_risk_text_rgbcolor>>'] = RGBColor(255,255,255)
                    
               
            index = 3
            for vunl in sorted_asc_vulns:
                doca.Tables(2).Cell(index, 1).Range.Text = vunl['<<level_risk>>']
                doca.Tables(2).Cell(index, 1).Shading.BackgroundPatternColor = vunl['<<level_risk_olecolor>>']
                doca.Tables(2).Cell(index, 1).Range.Font.Color = vunl['<<level_risk_text_olecolor>>']
                doca.Tables(2).Cell(index, 2).Range.Text =  "{0:3}. {1}".format(index-2,vunl['<<vulnerability_name>>']) 
                
                index = index + 1 
                doca.Tables(2).Rows.Add()
            doca.Tables(2).Cell(index, 1).Select() 
            wordapp.Selection.SelectRow() 
            wordapp.Selection.Cells.Delete(ShiftCells=win32.constants.wdDeleteCellsEntireRow)
            
            
            wordapp.Selection.HomeKey(Unit=win32.constants.wdStory)
            wordapp.Selection.Find.Execute('<<risk_list>>') 
            wordapp.Selection.Text = '\n'.join(risk_list)
            
            wordapp.Selection.HomeKey(Unit=win32.constants.wdStory)
            wordapp.Selection.Find.Execute('<<recomendation_list>>') 
            wordapp.Selection.Text = '\n'.join(remediation_list)
            
            doca.TablesOfContents(1).Update()
            wordapp.ActiveDocument.Save()
            doca.Close(SaveChanges=True)



            print("ERRORSOW")
            print("ERRORSOW")
            print("ERRORSOW")
            
            
            # SOW GENERATION
            concordancia_1 =  'a los sistemas' if int(data['<<no_targets>>']) > 1 else 'al sistema'
            concordancia_2 =  'de los portales' if int(data['<<no_targets>>']) > 1 else 'del portal'
            concordancia_3 =  'a los portales' if int(data['<<no_targets>>']) > 1 else 'al portal'

            
            Dict = dict({'<<Nombre_del_aplicativo_portada>>': str(data['<<name_app>>'] + ' - ' + data['<<analysis_version_format_01>>']),
             '<<Fecha_mes_y_año>>':data['<<date_format_02>>'], 
             '<<Folio>>':data['<<analysis_id>>'],
             '<<Fecha_ddmmaa_encabezado>>':data['<<date_format_01>>'],
             '<<Dirección_IP>>':data['<<scope_ip_01>>'],
             '<<request_folio>>':data['<<request_folio>>'],
             '<<Folio>>':data['<<analysis_id>>'],
             '<<analysis_version_format_02>>': data['<<analysis_version_format_02>>'],
             '<<Concordancia_1>>':concordancia_1,
             '<<Nombre_del_aplicativo_En_antecedentes>>':data['<<name_app>>'],
             '<<Nombre_del_servidor>>':data['<<scope_url_01>>'],
             '<<Nombre_del_aplicativo_Tabla>>':data['<<name_app>>'], 
             '<<Fechas_de_inicio>>': data['<<start_date>>'],
             '<<Fecha_Fin>>': data['<<finish_date>>'],
             '<<Fecha_tentativa_de_inicio>>': data['<<start_date_planned>>'],
             '<<Fecha_límite_para_la_actividad>>': data['<<due_date>>'], 
             '<<Concordancia_2>>': concordancia_2, 
             '<<URL_Acuerdos_tabla3>>': data['<<scope_url_01>>'],
             '<<Realiza_Firmas_de_aceptación>>': data['<<reviewer_01>>'],
             '<<Concordancia_3>>':concordancia_3})
            
            print("ERRORB")
            sow_template = os.path.join(dn,'templates','template-sre-sow-report.docx')
            sow_file_name = 'SOW - {}-{} {}'.format(data['<<analysis_id>>'],data['<<name_app>>'],data['<<analysis_version_format_01>>'])
            sow_full_file_name = os.path.join(dn,"outputs",sow_file_name+'.docx')
            doc = wordapp.Documents.Open(sow_template)
            doc.Activate()
            
            wordapp.Selection.GoTo(win32.constants.wdGoToPage, win32.constants.wdGoToAbsolute, "2")
            for From in Dict.keys():
                wordapp.ActiveWindow.ActivePane.View.SeekView =win32.constants.wdSeekMainDocument
                wordapp.Selection.Find.Execute(From, False, False, False, False, False, True, win32.constants.wdFindContinue, False, Dict[From], win32.constants.wdReplaceAll) 
                wordapp.ActiveWindow.ActivePane.View.SeekView = win32.constants.wdSeekCurrentPageHeader
                wordapp.Selection.Find.Execute(From, False, False, False, False, False, True, win32.constants.wdFindContinue, False, Dict[From], win32.constants.wdReplaceAll)     
            
            wordapp.Selection.GoTo(win32.constants.wdGoToPage, win32.constants.wdGoToAbsolute, "1")
            wordapp.ActiveDocument.SaveAs(sow_full_file_name)
            doc.Close(False)
            
            try:
               covx_to_pdf(full_vulns_table_file_name, full_vulns_table_file_name.replace('.docx', '.pdf'), wordapp)
               covx_to_pdf(os.path.join(dn,"outputs",name_file), os.path.join(dn,"outputs",name_file).replace('.docx', '.pdf'),wordapp)
               covx_to_pdf(sow_full_file_name, sow_full_file_name.replace('.docx', '.pdf'),wordapp)
            except Exception as e: 
               print(e) 
            
            
            for f in vulnerabilities_tables:
                # Delete temporal files generated
                if os.path.isfile(f): # this makes the code more robust
                        os.remove(f)
            
            if os.path.isfile(vulns_table_file_name): # this makes the code more robust
                        os.remove(vulns_table_file_name)

        except Exception as e: 
           print(e) 
   else:
        
        # SOW GENERATION
        concordancia_1 =  'a los sistemas' if int(data['<<no_targets>>']) > 1 else 'al sistema'
        concordancia_2 =  'de los portales' if int(data['<<no_targets>>']) > 1 else 'del portal'
        concordancia_3 =  'a los portales' if int(data['<<no_targets>>']) > 1 else 'al portal'
        
        Dict = dict({'<<Nombre_del_aplicativo_portada>>': str(data['<<name_app>>'] + ' - ' + data['<<analysis_version_format_01>>']),
             '<<Fecha_mes_y_año>>':data['<<date_format_02>>'], 
             '<<Folio>>':data['<<analysis_id>>'],
             '<<Fecha_ddmmaa_encabezado>>':data['<<date_format_01>>'],
             '<<Dirección_IP>>':data['<<scope_ip_01>>'],
             '<<request_folio>>':data['<<request_folio>>'],
             '<<Folio>>':data['<<analysis_id>>'],
             '<<Concordancia_1>>':concordancia_1,
             '<<Nombre_del_aplicativo_En_antecedentes>>':data['<<name_app>>'],
             '<<Nombre_del_servidor>>':data['<<scope_url_01>>'],
             '<<Nombre_del_aplicativo_Tabla>>':data['<<name_app>>'], 
             '<<Fechas_de_inicio>>': data['<<start_date>>'],
             '<<Fecha_Fin>>': data['<<finish_date>>'],
             '<<Fecha_tentativa_de_inicio>>': data['<<start_date_planned>>'],
             '<<Fecha_límite_para_la_actividad>>': data['<<due_date>>'], 
             '<<Concordancia_2>>': concordancia_2, 
             '<<URL_Acuerdos_tabla3>>': data['<<scope_url_01>>'],
             '<<Realiza_Firmas_de_aceptación>>': data['<<reviewer_01>>'],
             '<<Concordancia_3>>':concordancia_3})

        sow_template = os.path.join(dn,'templates','template-sre-sow-report.docx')
        sow_file_name = 'SOW - {}-{} {}'.format(data['<<analysis_id>>'],data['<<name_app>>'],data['<<analysis_version_format_01>>'])
        sow_full_file_name = os.path.join(dn,"outputs",sow_file_name+'.docx')
        wordapp = win32.gencache.EnsureDispatch("Word.Application")
        wordapp.Visible = visible_mode_win32com
        wordapp.DisplayAlerts = False
        doc = wordapp.Documents.Open(sow_template)
        doc.Activate()
        
        wordapp.Selection.GoTo(win32.constants.wdGoToPage, win32.constants.wdGoToAbsolute, "2")
        for From in Dict.keys():
            
            wordapp.ActiveWindow.ActivePane.View.SeekView =win32.constants.wdSeekMainDocument
            wordapp.Selection.Find.Execute(From, False, False, False, False, False, True, win32.constants.wdFindContinue, False, Dict[From], win32.constants.wdReplaceAll) 
            wordapp.ActiveWindow.ActivePane.View.SeekView = win32.constants.wdSeekCurrentPageHeader
            wordapp.Selection.Find.Execute(From, False, False, False, False, False, True, win32.constants.wdFindContinue, False, Dict[From], win32.constants.wdReplaceAll)     
        
        wordapp.Selection.GoTo(win32.constants.wdGoToPage, win32.constants.wdGoToAbsolute, "1")
        wordapp.ActiveDocument.SaveAs(sow_full_file_name)
        doc.Close(False)
        
        try:
           covx_to_pdf(os.path.join(dn,"outputs",name_file), os.path.join(dn,"outputs",name_file).replace('.docx', '.pdf'),wordapp)
           covx_to_pdf(sow_full_file_name, sow_full_file_name.replace('.docx', '.pdf'),wordapp)
        except Exception as e: 
           print(e) 

   wordapp.Application.Quit() 

def markdown_tojson(markdown_filename):
   
   # Create target directory & all intermediate directories if don't exists
   if not os.path.isfile(markdown_filename):
       print('No existe ese archivo.')
       exit()
   
   markdown_file  = open(markdown_filename,mode="r", encoding="utf-8") 
   markdown_file_data = markdown_file.read()
   markdown_file.close()
   html_start = '<html><body>'
   html_end = "</body></html>"
   html_str = mistune.markdown(markdown_file_data)
   complete_html = html_start+html_str+html_end
   
   Html_file= open(os.path.join(dn,"outputs",'temp.html') ,"w")
   Html_file.write(complete_html)
   Html_file.close()
   
   soup = BeautifulSoup(complete_html, "lxml")
   
   invalid_tags = ['b','a', 'i', 'u']
   
   for tag in invalid_tags: 
       for match in soup.findAll(tag):
           match.replaceWithChildren()
   
   with open(os.path.join(dn,"outputs",'temp-nohref.html'), "w") as file:
       file.write(str(soup))
   
   relevant_document_html_tags = ['h1','h2','h3','h4','h5','h6','p']
   vuln_tags = ['h2','h6','p']
   data_structured = {}
   
   
   for element in soup.find_all('h1'):  
       if  element.text == 'App information':  
           tag = element.find_next_sibling(relevant_document_html_tags)
           while tag.name == 'h6' or tag.name == 'p':
               content = "" 
               if tag.name == 'h6': 
                   key = tag.text                
               elif tag.name == 'p':
                   subtag = tag                
                   while subtag.name == 'p':
                       content = content + subtag.text     
                       subtag = subtag.find_next_sibling(relevant_document_html_tags)
               data_structured["<<"+key+">>"] = content
               tag = tag.find_next_sibling(relevant_document_html_tags)
   
   
   data_structured['<<vulnerabilities>>'] = []
   vuln = {}  
   key = ""
   content = ""
   for element in soup.find_all('h1',text ='Vulnerabilities'):
           subelement = element.find_next_sibling()  
           while subelement is not None:
               if subelement.name == 'h1': 
                   break # Stop list all elements final state
               elif subelement.name == 'h2':
                   if not vuln:
                       vuln["<<vulnerability_name>>"] = subelement.text               
                       vuln["<<vulnerability_evidences>>"] = []
                       subelement = subelement.find_next_sibling()
                   else:
                       vuln["<<"+key+">>"] = content   
                       data_structured['<<vulnerabilities>>'].append(vuln)
                       print()
                       print("New1 vuln appended "+ str(vuln)) 
                       vuln = {} 
               elif subelement.name == 'h5':
                   subelement = subelement.find_next_sibling('h2') # Go to next vuln
                   if subelement is None:
                       break
               elif subelement.name == 'h6':                
                   print("actual key value: "+ key)
                   print("actual content value: "+ content)
                   print("next element current subelement: "+ subelement.text)
                   if not key:
                       key = subelement.text
                       subelement = subelement.find_next_sibling() 
                   else:
                       vuln["<<"+key+">>"] = content                   
                       key = subelement.text
                       content = ""
                       subelement = subelement.find_next_sibling()                        
               elif subelement.name == 'p':
                   if subelement.find_next_sibling().name == 'p':
                       content = content + subelement.text + '\r' 
                       print("1content value: "+ content)   
                       subelement = subelement.find_next_sibling() 
                   else:  
                       content = content + subelement.text
                       print("2content value: "+ content)   
                       subelement = subelement.find_next_sibling()
           vuln["<<"+key+">>"] = content            
           data_structured['<<vulnerabilities>>'].append(vuln) 
           print()
           print("New2 vuln appended "+ str(vuln)) 
   
  
   count = 0
   count_vuln = 0
   for element in soup.find_all('h2'): 
       if element.find_previous_sibling('h1').text == "Vulnerabilities":  
           print("Vulnerability: {}".format(element.text))
           next_element = element.find_next_sibling('h5')
           vulnerability_evidences = [] 
           
           while next_element.name == 'h5':
               evidence = {}
               vulnerability_evidence_note = "" 
               vulnerability_evidence_image_path = "" 
               content = ""
               print("Evidence {}: {}".format(count,next_element.text))
               evidence_element = next_element.find_next_sibling(['h6','p'])
               
               while evidence_element.name == 'h6' or evidence_element.name == 'h2':
                   if  evidence_element.name == 'h2':
                       break
                   elif  evidence_element.name == 'h6' and evidence_element.text == "vulnerability_evidence_image_path":
                       sub_tag_p = evidence_element.find_next_sibling(['h6','p'])
                       content = ""
                       while sub_tag_p.name == 'p':
                           content = content + sub_tag_p.text + '\r'     
                           if sub_tag_p.find_next_sibling(['h6','p']) is None:
                               break
                           else:
                               sub_tag_p = sub_tag_p.find_next_sibling(['h6','p'])  
                       evidence["<<vulnerability_evidence_image_path>>"] = content
                   elif evidence_element.name == 'h6' and evidence_element.text == "vulnerability_evidence_note":
                       sub_tag_p = evidence_element.find_next_sibling(['h6','p'])
                       content = ""
                       while sub_tag_p.name == 'p':
                           content = content + sub_tag_p.text
                           if sub_tag_p.find_next_sibling(['h6','p']) is None:
                               break
                           else:
                               sub_tag_p = sub_tag_p.find_next_sibling(['h6','p']) 
                       evidence["<<vulnerability_evidence_note>>"] = content
                   if  evidence_element.name == 'h5' and evidence_element.text == "evidence":
                       print("Agregando evidencia a la misma vulnerabilidad")
                       vulnerability_evidences.append(evidence)
                   
                   if evidence_element.find_next_sibling(['h6','h2','h5']) is not None:                
                       evidence_element = evidence_element.find_next_sibling(['h6','h2','h5']) 
                   else:
                       break
               
               vulnerability_evidences.append(evidence)
               print("Evidence conent {}: {}".format(evidence["<<vulnerability_evidence_image_path>>"],evidence["<<vulnerability_evidence_note>>"]))
               
               if next_element.find_next_sibling(['h5','h2']) is not None: 
                   count = count + 1               
                   next_element = next_element.find_next_sibling(['h5','h2']) 
               else:
                   break
           data_structured['<<vulnerabilities>>'][count_vuln]["<<vulnerability_evidences>>"] = vulnerability_evidences
           count_vuln = count_vuln + 1
           
   json_data = json.dumps(data_structured, indent=2, separators=(',', ':'))
   
   return json_data

def removeLeadingZeros(str):
 
    # Regex to remove leading
    # zeros from a string
    regex = "^0+(?!$)"
 
    # Replaces the matched
    # value with given string
    str = re.sub(regex, "", str)
 
    print(str)


def decode_base64(data, altchars=b'+/'):
    """Decode base64, padding being optional.

    :param data: Base64 data as an ASCII byte string
    :returns: The decoded byte string.

    """
    data = re.sub(rb'[^a-zA-Z0-9%s]+' % altchars, b'', data)  # normalize
    missing_padding = len(data) % 4
    if missing_padding:
        data += b'='* (4 - missing_padding)
    return base64.b64decode(data, altchars)



@wda.get('/wda/')
async def find_all_wdas():
    return serializeList(db.local.wda.find())

@wda.get('/wdas/')
async def find_all_wdas_id():
    print(db.local.wda.find())
    return serializeList(db.local.wda.find())

@wda.get('/wda/{id}')
async def find_one_wda(id):
    html_content = """
    <html>
        <head>
            <title>Not found</title>
        </head>
        <body>
            <h1>Sorry not found</h1>
        </body>
    </html>
    """

    
    
    json_data = db.local.wda.find_one({"_id":ObjectId(id)})

    if '_id' in json_data:
                del json_data['_id']

    outputs_directory = os.path.join(dn,"outputs")

    for i in range(len(json_data["<<vulnerabilities>>"])):
        count = 1
        for k in range(len(json_data["<<vulnerabilities>>"][i]["<<vulnerability_evidences>>"])):
            image = os.path.join(outputs_directory,"vulnerability_{}.png".format(str(count)))
            print(os.path.join(outputs_directory,"vulnerability_{}.png".format(str(count))))
            fh = open(image, "wb")
            fh.write(base64.b64decode(json_data["<<vulnerabilities>>"][i]["<<vulnerability_evidences>>"][k]["<<vulnerability_evidence_image_path>>"]))
            fh.close()
            json_data["<<vulnerabilities>>"][i]["<<vulnerability_evidences>>"][k]["<<vulnerability_evidence_image_path>>"] = image            

  
    generate_report(json_data)
    zip_name = os.path.join(dn,"report")
    zip_path = os.path.join(dn,"report.zip")

    if os.path.exists(zip_path):
        os.remove(zip_path)
    else:
        print("Can not delete the file as it doesn't exists")
 
    
    files_in_directory = os.listdir(outputs_directory)
    shutil.make_archive(zip_name, 'zip', outputs_directory)


    for files in os.listdir(outputs_directory):
        path = os.path.join(outputs_directory, files)
        try:
            shutil.rmtree(path)
        except OSError:
            os.remove(path)

    try: 
        response = FileResponse(path=zip_path, filename="report.zip", media_type='application/zip')    
        return response    
    except IOError:
    # handle file not exist case here
        return HTMLResponse(content=html_content, status_code=200)


@wda.post('/wda/')
async def create_wda(request: Request):
    
    #wda_dict = wdaDict(wda)
    
    #print(wda_dict)

    response = "{}"

    json_request = await request.json()
    
    print("Validating report JSON...")
    if '<<template_name>>' not in json_request:
        print("The file has the correct report format.")
        print("Insertando datos en base de datos...")
        new_report = json_request 
        db.insert(new_report)
        response = json_request
    else:
        print("The file is not in the correct report format!!!")

    return response

@wda.put('/wda/{id}')
async def update_wda(id,wda: VulnerabilityAnalisis):
    db.local.wda.find_one_and_update({"_id":ObjectId(id)},{
        "$set":dict(wda)
    })
    return serializeDict(db.local.wda.find_one({"_id":ObjectId(id)}))

@wda.delete('/wda/{id}')
async def delete_wda(id,wda: VulnerabilityAnalisis):
    return serializeDict(db.local.wda.find_one_and_delete({"_id":ObjectId(id)}))